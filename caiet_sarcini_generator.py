# -*- coding: utf-8 -*-
"""
Generator CAIET DE SARCINI instalaţii electrice (.docx) — livrabil distinct de memoriu.

Memoriul = CE s-a proiectat (soluţie, breviar). Caietul = CUM se execută (materiale,
pozare, verificări, recepţie, sarcini executant/beneficiar) — pentru firma de execuţie.

Emis DOAR la fazele care includ PT (DTAC+PT / PT) — gate-ul e în endpoint (_is_pt).
Structura validată de Dan (2026-07-24): 7 capitole pe TEME (model INSTAUDITOR) +
livrare/depozitare/manipulare (model Hunedoara), adaptate REZIDENŢIAL: cabluri CYY-F,
fără referinţe la alte birouri, ~10-14 pagini.

Dinamice: 1.4 nominalizări planşe (payload / plansa_numbering la PT), menţiunea FV
(solar + FV_PACKAGES), cap. 4 referinţa la planşa de forţă, cap. 7 specificaţia
tablourilor din circuitele reale (TEG / TES n / TE-CT).
"""
import io
import re

from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Infrastructura DOCX refolosită din memoriu (stiluri, copertă, helpers, gate PT).
from memoriu_generator import (
    _setup_document, _page_coperta, _add_heading, _add_para,
    _set_run_font, _set_table_fixed, set_table_borders, _is_pt,
    _cs_din_circuite,          # inventarul de curenti slabi, insumat de pe circuite
    _alim_din_firida,          # alimentare din firida blocului (vs bransament propriu)
)

# =============================================================================
# TEXT FIX — capitolele statice (redactate din modelele validate, adaptate
# rezidenţial: CYY-F, fără alte birouri, fără curenţi slabi/anti-efracţie).
# Tuple: (kind, text); kind in {"h1","h2","p","li"}
# =============================================================================

_CS_CAP1 = [
    ("h1", "1. GENERALITĂŢI"),

    ("h2", "1.1. Obiectul şi domeniul de aplicare"),
    ("p", "Prezentul caiet de sarcini se referă la executarea instalaţiilor electrice interioare "
          "aferente obiectivului menţionat în foaia de capăt: instalaţii de iluminat, instalaţii de "
          "prize şi instalaţii de forţă (alimentări receptoare dedicate), inclusiv instalaţia de "
          "legare la pământ."),
    ("p", "Instalaţiile electrice de utilizare se vor executa numai de către firme atestate şi/sau "
          "electricieni autorizaţi conform ordinelor ANRE, având gradul de competenţă corespunzător "
          "lucrării executate. Instalaţiile electrice se vor executa cu respectarea normelor şi "
          "reglementărilor în vigoare şi având avizul tehnic de racordare al furnizorului de energie "
          "electrică. Lucrările referitoare la branşamentul obiectivului sunt cuprinse în proiectul "
          "elaborat de furnizorul de energie electrică."),   # ancora _ALIM_CAIET (alimentare din firidă)
    ("p", "Antreprenorul are obligaţia de a executa lucrările conform proiectului, condiţiilor "
          "contractuale şi prescripţiilor tehnice în vigoare. În timpul execuţiei, orice modificări "
          "sau completări ale proiectului se fac numai cu respectarea dispoziţiilor legale şi cu "
          "acordul scris al proiectantului şi al beneficiarului."),
    ("p", "Cerinţele prezentului caiet de sarcini nu îl exonerează pe antreprenor de "
          "responsabilitatea de a realiza şi alte verificări, încercări şi activităţi pe care le "
          "consideră necesare pentru asigurarea calităţii execuţiei şi a materialelor."),

    ("h2", "1.2. Materiale — condiţii tehnice de calitate"),
    ("p", "Pentru executarea circuitelor din prezentul proiect se folosesc numai materiale "
          "omologate, noi, însoţite de certificate de calitate şi agremente tehnice, şi anume:"),
    ("li", "tuburi de protecţie IPEY (montaj îngropat) cu mufe şi curbe aferente; se vor folosi "
           "numai tuburi pentru care există piese de îmbinare uzinate; tuburile montate aparent vor "
           "fi incombustibile sau greu combustibile, cu degajări reduse de gaze;"),
    ("li", "cabluri din cupru tip CYY-F, cu conductor de cupru conform SR CEI 60228, izolaţie şi "
           "manta din PVC, întârziere la propagarea flăcării, tensiune nominală 0,6/1 kV, "
           "temperatură maximă de funcţionare 70°C pe conductor; secţiunile şi numărul de "
           "conductoare — conform schemelor monofilare şi listei de cantităţi din proiect;"),
    ("li", "corpuri de iluminat cu sursă LED, conform SR EN 60598-1, alese potrivit destinaţiei "
           "fiecărei încăperi; în spaţiile umede sau exterioare — grad de protecţie minim IP44 "
           "conform SR EN 60529;"),
    ("li", "aparataj terminal (întreruptoare, comutatoare, prize cu contact de protecţie) pentru "
           "montaj îngropat în doze de aparat; întreruptoarele pentru circuitele de iluminat vor "
           "avea curentul nominal de 10 A, iar prizele — 16 A, obligatoriu cu contact de protecţie;"),
    ("li", "tablouri electrice modulare, echipate conform schemelor monofilare din proiect, "
           "executate de firme specializate; aparatele de protecţie conform SR EN 60898-1 "
           "(întreruptoare automate), respectiv SR EN 61008-1 / SR EN 61009-1 (dispozitive "
           "diferenţiale)."),
    ("p", "Materialele, aparatele şi echipamentele ale căror caracteristici nu corespund cu cele "
          "din proiect sau care prezintă defecte de calitate vor fi respinse."),

    ("h2", "1.3. Livrare, depozitare, manipulare"),
    ("p", "Tuburi de protecţie: livrarea se face în loturi compacte, pe tipodimensiuni, cu etichete "
          "pe care se specifică clar tipul şi cantitatea. Depozitarea se face în locuri ferite de "
          "umezeală şi temperaturi negative, fără expunere la radiaţii solare, de preferinţă în "
          "spaţii supravegheate. Manipularea se face cu grijă, de preferinţă manual, pentru a nu se "
          "produce deteriorări ale tuburilor sau ale accesoriilor acestora."),
    ("p", "Cabluri electrice: livrarea se face pe tamburi sau în colaci, pe tipodimensiuni, cu "
          "etichete pe care se specifică clar tipul şi cantitatea. Depozitarea se face în locuri "
          "ferite de umezeală şi temperaturi negative. Manipularea se face cu grijă; rostogolirea "
          "tamburului este admisă doar în sensul indicat pe tambur, pentru a nu se produce "
          "desfăşurarea, îndoirea sau ruperea spirelor."),
    ("p", "Corpuri de iluminat: livrarea se face în ambalajele originale, inclusiv pentru "
          "accesorii, etichetate corespunzător şi însoţite de documentele de livrare (factură, "
          "certificat de calitate, certificat de garanţie). Depozitarea se face obligatoriu în "
          "spaţii închise, ferite de umezeală, fără expunere la intemperii şi fără materiale "
          "inflamabile în apropiere. Manipularea se face numai manual, respectând indicaţiile de pe "
          "ambalaj (FRAGIL, A NU SE RĂSTURNA, A SE FERI DE UMEZEALĂ)."),
    ("p", "Tablouri electrice: livrarea se face cu tabloul ambalat în folie de protecţie, însoţit "
          "de certificatul de calitate şi de garanţie emis de producător. Transportul şi "
          "manipularea se fac fără deteriorări mecanice, în poziţie verticală, iar depozitarea în "
          "spaţii închise, fără umiditate şi fără gaze corozive (temperaturi 0–40°C)."),
]

_CS_SARCINI = [
    ("h2", "1.5. Sarcini pentru executant"),
    ("p", "Pentru realizarea în bune condiţii a tuturor lucrărilor din prezentul proiect, "
          "executantul va desfăşura următoarele activităţi:"),
    ("li", "studierea proiectului pe baza pieselor scrise şi desenate din documentaţie, precum şi a "
           "legislaţiei, standardelor şi instrucţiunilor tehnice, astfel ca până la începerea "
           "execuţiei să fie clarificate toate lucrările ce urmează a fi executate;"),
    ("li", "sesizarea proiectantului, în termen legal, asupra eventualelor neconcordanţe între "
           "elementele grafice şi cifrice, în vederea rezolvării lor;"),
    ("li", "asigurarea aprovizionării cu materialele şi produsele din proiect;"),
    ("li", "sesizarea proiectantului în cazul imposibilităţii procurării unor materiale sau "
           "aparataje prevăzute în documentaţie, prezentând în acelaşi timp o ofertă de material "
           "similar, cu caracteristici cel puţin echivalente din punct de vedere tehnic;"),
    ("li", "asigurarea forţei de muncă şi a mijloacelor de mecanizare, ritmic, în concordanţă cu "
           "graficul de execuţie şi cu termenele parţiale sau finale stabilite;"),
    ("li", "respectarea cu stricteţe a tehnologiei de lucru caracteristice."),
    ("p", "Executantul este obligat să păstreze pe şantier, la punctul de lucru, pe toată perioada "
          "de execuţie şi a efectuării probelor, întreaga documentaţie pe baza căreia se execută "
          "lucrările, inclusiv dispoziţiile de şantier date pe parcurs. Această documentaţie, "
          "împreună cu procesele-verbale de lucrări ascunse şi documentele care atestă calitatea "
          "materialelor, va fi pusă la dispoziţia organelor de control (Inspecţia de Stat în "
          "Construcţii). Modificările prevederilor documentaţiei tehnice se vor executa numai cu "
          "avizul scris al proiectantului şi vor fi stipulate şi în partea desenată a "
          "documentaţiei."),

    ("h2", "1.6. Sarcini pentru beneficiar"),
    ("p", "Beneficiarului, prin dirigintele de şantier, îi revin următoarele sarcini:"),
    ("li", "recepţionează documentaţia primită de la proiectant, verificând piesele scrise şi "
           "desenate, corelarea între ele şi exactitatea elementelor (lungimi, trasee etc.);"),
    ("li", "sesizează proiectantul asupra neconcordanţelor sau situaţiilor specifice apărute în "
           "execuţie, în scopul analizei comune şi găsirii rezolvării urgente;"),
    ("li", "anunţă proiectantul în vederea prezentării la fazele determinante, la punerea în "
           "funcţiune sau în alte situaţii care impun schimbarea soluţiilor din proiect;"),
    ("li", "nu acceptă la montaj modificări faţă de documentaţie decât cu avizul proiectantului;"),
    ("li", "urmăreşte ritmic execuţia lucrărilor în scopul respectării documentaţiei, controlând "
           "calitatea lucrărilor şi participând la confirmarea lucrărilor ascunse şi a cantităţilor "
           "de lucrări efectuate de executant, la nivelul fiecărei faze determinante;"),
    ("li", "nu acceptă trecerea la o altă fază sau recepţia lucrărilor executate fără atestarea "
           "tuturor elementelor care concură la o bună calitate a materialelor şi a execuţiei."),
]

# Cap. 2 — lista de normative VALIDATĂ de Dan (curăţată, doar în vigoare).
_CS_NORMATIVE = [
    ("h1", "2. NORMATIVE, PRESCRIPŢII ŞI STANDARDE DE REFERINŢĂ"),
    ("p", "Elaborarea prezentei documentaţii tehnice s-a făcut în conformitate cu prevederile "
          "normativelor, prescripţiilor tehnice şi standardelor în vigoare. Constructorul va avea "
          "în vedere ca toate materialele şi echipamentele puse în operă să fie conforme cu "
          "cerinţele specificate în următoarele:"),
    ("h2", "Normative"),
    ("li", "I7-2011 — Normativ pentru proiectarea, execuţia şi exploatarea instalaţiilor electrice "
           "aferente clădirilor;"),
    ("li", "NP 061-2002 — Normativ pentru proiectarea şi executarea sistemelor de iluminat "
           "artificial din clădiri;"),
    ("li", "NTE 007/08/00 — Normativ pentru proiectarea şi executarea reţelelor de cabluri "
           "electrice;"),
    ("li", "PE 132-2003 — Normativ pentru proiectarea reţelelor electrice de distribuţie publică;"),
    ("li", "I18/1-2001 — Instalaţii electrice interioare de curenţi slabi;"),
    ("li", "I18/2-2002 — Instalaţii de semnalizare a incendiilor;"),
    ("li", "C 56-2002 — Verificarea calităţii şi recepţia lucrărilor de instalaţii."),
    ("h2", "Standarde"),
    ("li", "SR EN 60617 — Simboluri grafice pentru scheme electrice;"),
    ("li", "SR EN 60898-1 — Întreruptoare automate pentru instalaţii casnice;"),
    ("li", "SR EN 61008-1 / SR EN 61009-1 — Dispozitive de protecţie la curent diferenţial "
           "rezidual (RCCB / RCBO);"),
    ("li", "SR EN 60529 — Grade de protecţie asigurate prin carcase (Cod IP);"),
    ("li", "SR EN 60598-1 — Corpuri de iluminat;"),
    ("li", "SR HD 60364 (seria) — Instalaţii electrice de joasă tensiune;"),
    ("li", "SR EN 62305 (seria) — Protecţia împotriva trăsnetului;"),
    ("li", "SR EN 50522 — Legarea la pământ a instalaţiilor de energie."),
    ("h2", "Legislaţie"),
    ("li", "Legea 10/1995 — privind calitatea în construcţii (republicată);"),
    ("li", "HG 766/1997 — pentru aprobarea unor regulamente privind calitatea în construcţii."),
]

_CS_NORMATIVE_FV = [
    ("h2", "Sistem fotovoltaic"),
    ("li", "SR EN 62446 — Sisteme fotovoltaice conectate la reţea;"),
    ("li", "IEC 62548 — Cerinţe de proiectare pentru generatoare fotovoltaice;"),
    ("li", "Ordinele ANRE privind prosumatorii."),
]

_CS_CAP3 = [
    ("h1", "3. EXECUTAREA INSTALAŢIILOR ELECTRICE"),
    ("p", "Înainte de a începe montarea instalaţiilor electrice se vor verifica şi identifica "
          "viitoarele trasee electrice. La traseele alese se va verifica dacă: lungimea traseelor "
          "este cea mai scurtă; s-au respectat distanţele minime admise până la elementele altor "
          "instalaţii şi faţă de elementele de construcţie combustibile; s-au evitat locurile "
          "periculoase. Toate traseele care nu satisfac condiţiile impuse vor fi reexaminate şi "
          "retrasate. Conductoarele electrice şi tuburile de protecţie se amplasează faţă de "
          "conductele altor instalaţii şi faţă de elementele de construcţie respectând distanţele "
          "minime din tabelul 3.1, art. 3.5 din Normativul I7-2011."),
    ("p", "Este strict interzis a se executa de către instalatori străpungeri sau goluri prin "
          "structura de rezistenţă a construcţiei. Acestea se admit numai pe baza unui aviz scris "
          "de la proiectantul structurii de rezistenţă şi se execută numai cu mijloace mecanizate "
          "adecvate."),
    ("h2", "Tuburi de protecţie şi cabluri"),
    ("p", "Tuburile izolante prevăzute pentru circuitele de iluminat şi prize se montează îngropat "
          "în elementele de construcţie. Traseele tuburilor vor fi verticale prin pereţi şi "
          "orizontale şi rectilinii prin planşee, pozarea făcându-se pe distanţa cea mai scurtă. La "
          "schimbările de direcţie, raza de curbură va fi conform prescripţiilor furnizorului, dar "
          "minimum 5D (D — diametrul tubului). Fixarea tuburilor de elementele de construcţie din "
          "BCA sau cărămidă se face cu ipsos, din 0,5 în 0,5 m."),
    ("p", "Circuitele se execută cu cabluri din cupru tip CYY-F, protejate în tuburi IPEY, cu "
          "secţiunile din schemele monofilare. Tragerea conductelor în tuburi se va executa numai "
          "după montarea tuburilor (la montaj îngropat — după uscarea tencuielilor). Trecerea "
          "cablurilor prin pereţi şi planşee se face în tuburi de protecţie; trecerea prin peretele "
          "exterior va fi perfect etanşă, pentru a preveni infiltraţiile. La conductoare şi cabluri "
          "se va verifica continuitatea electrică pe fiecare colac/tambur înainte de montare, cu "
          "ohmetrul; cele întrerupte vor fi respinse."),
    ("p", "Legăturile electrice între conductoarele de cupru se execută numai în doze, prin "
          "cleme corespunzătoare secţiunii. Se interzice utilizarea cordoanelor flexibile pentru "
          "executarea instalaţiilor electrice fixe. La legarea la aparate se prevăd lungimi "
          "suplimentare de circa 5–10% pentru evitarea solicitării conductorului. Calitatea "
          "circuitelor electrice se verifică după tragerea conductelor în tuburi, înainte de "
          "acoperire; rezistenţa de izolaţie între conducte şi pământ, măsurată cu circuitul "
          "deconectat, trebuie să fie de cel puţin 0,5 MΩ."),
    ("h2", "Tablouri electrice"),
    ("p", "Tablourile electrice se comandă pentru execuţie la furnizori specializaţi. Tablourile "
          "se instalează astfel încât înălţimea laturii de sus să nu depăşească 2,3 m, se montează "
          "vertical şi se fixează cu dibluri şi şuruburi, pentru a nu fi supuse vibraţiilor sau "
          "deplasărilor. Toate circuitele vor fi prevăzute cu inscripţii vizibile şi neechivoce "
          "care să indice destinaţia fiecărui circuit, amplasate cu vedere din direcţia de "
          "deservire a tabloului."),
    ("p", "Înainte de racordarea circuitelor la tablouri se verifică: integritatea construcţiei, "
          "existenţa şi integritatea etichetelor şi a aparatelor, strângerea legăturilor, fixarea "
          "aparatelor şi legătura de protecţie la bara PE. La verificarea instalării tablourilor se "
          "controlează modul şi calitatea fixării, înălţimile de montaj, distanţele admise faţă de "
          "elementele altor instalaţii, existenţa tuturor aparatelor de protecţie prevăzute în "
          "proiect şi calitatea legăturilor."),
    ("h2", "Aparate şi corpuri de iluminat"),
    ("p", "Dozele de aparat se montează îngropat. Montarea aparatelor se face în ultima fază de "
          "execuţie a finisajelor, după zugrăveli şi vopsitorii. Întreruptoarele şi comutatoarele "
          "se montează astfel încât să întrerupă faza la corpul de iluminat, la înălţimile din "
          "planşele proiectului. Prizele vor fi obligatoriu cu contact de protecţie, conectarea "
          "conductorului de protecţie la borna corespunzătoare fiind obligatorie; prizele trebuie "
          "să reziste tensiunii mecanice exercitate la tragerea ştecherului fără a fi ţinute cu "
          "mâna."),
    ("p", "Corpurile de iluminat se aleg şi se montează respectând prevederile NP 061. Conductorul "
          "de fază se leagă în dulia lămpii la borna din interior, iar conductorul de nul la partea "
          "filetată a duliei. Se interzice suspendarea corpurilor de iluminat direct de conductele "
          "de alimentare; dispozitivele de suspendare se aleg astfel încât să suporte fără "
          "deformări o greutate egală cu de 5 ori greutatea corpului de iluminat, minimum 10 kg."),
    # ILUMINAT DE SIGURANTA (I7-2011 cap. 7.23 + SR EN 1838). Text STATIC, ca restul caietului:
    # cerintele de executie sunt aceleasi indiferent daca proiectul are sau nu corpuri de siguranta
    # (caietul descrie CUM se executa, memoriul descrie CE s-a proiectat).
    ("p", "Corpurile de iluminat de siguranţă (evacuare şi antipanică) vor fi executate din "
          "materiale cu clasa B de reacţie la foc şi se montează la minimum 2 m faţă de pardoseală, "
          "pe căile de evacuare şi deasupra ieşirilor, la o distanţă de cel mult 15 m între ele. "
          "Corpurile autonome şi kiturile de emergenţă vor avea autonomie de minimum 2 ore şi "
          "punere în funcţiune în cel mult 5 secunde de la dispariţia tensiunii normale."),
    ("p", "Circuitul de alimentare a corpurilor de evacuare este dedicat, fără protecţie "
          "diferenţială şi fără întreruptor pe traseu, astfel încât să nu poată fi scos din "
          "funcţiune odată cu circuitele normale de iluminat. La recepţie se verifică aprinderea "
          "corpurilor la întreruperea alimentării normale şi menţinerea lor în funcţiune pe durata "
          "de autonomie declarată. Verificarea se repetă periodic pe durata exploatării, iar "
          "rezultatele se consemnează în registrul de întreţinere al instalaţiei."),
]

_CS_CAP4_PRINCIPII = [
    ("p", "Instalaţia de protecţie prin legare la pământ se realizează pentru prevenirea "
          "accidentelor produse prin atingere indirectă. La această instalaţie se racordează toate "
          "elementele conductive care nu fac parte din circuitele curenţilor de lucru, dar care "
          "accidental ar putea intra sub tensiune: carcasele şi elementele de susţinere metalice "
          "ale echipamentelor electrice, barele PE ale tablourilor electrice, contactele de "
          "protecţie ale prizelor."),
    ("p", "Platbanda se pozează pe conturul fundaţiei, înglobată în betonul fundaţiei, astfel "
          "încât să fie acoperită cu un strat de beton de cel puţin 3 cm. Asigurarea continuităţii "
          "electrice a legăturilor se face prin îmbinări sudate de bună calitate; legătura dintre "
          "priza de pământ şi bara PE a tabloului se realizează prin piesa de separaţie, care "
          "permite măsurarea rezistenţei de dispersie."),
    ("p", "Rezistenţa de dispersie a prizei de pământ trebuie să fie de cel mult 4 Ω (respectiv "
          "1 Ω în cazul prizei comune cu instalaţia de protecţie împotriva trăsnetului). Dacă la "
          "măsurătorile efectuate rezistenţa de dispersie rezultă mai mare decât cea prescrisă, se "
          "vor lua măsuri de îmbunătăţire prin electrozi suplimentari (platbandă OL-Zn 40×4 mm sau "
          "electrozi verticali), până la încadrarea în valorile admise."),
    ("p", "Schema de legare la pământ este TN-S: conductorul de protecţie PE este distinct de "
          "conductorul neutru N pe întreaga instalaţie, iar toate masele instalaţiei electrice "
          "sunt legate la PE."),
]

_CS_CAP5 = [
    ("h1", "5. VERIFICĂRI, PROBE ŞI RECEPŢIA LUCRĂRILOR"),
    ("p", "În timpul execuţiei se face o verificare preliminară, iar după executarea instalaţiei — "
          "verificarea definitivă, înainte de punerea în funcţiune, conform C 56-2002 şi I7-2011."),
    ("p", "Verificarea preliminară presupune:"),
    ("li", "verificarea înainte de montaj a calităţii materialelor şi a continuităţii electrice a "
           "conductoarelor;"),
    ("li", "verificarea după montaj a continuităţii electrice a instalaţiei, înaintea acoperirii "
           "de orice fel;"),
    ("li", "verificarea calităţii tuburilor montate în elementele de construcţie;"),
    ("li", "verificarea aparatelor electrice."),
    ("p", "Verificarea definitivă cuprinde verificări prin examinări vizuale şi verificări prin "
          "încercări. Prin examinare vizuală se stabileşte dacă: au fost aplicate măsurile de "
          "protecţie împotriva şocurilor electrice; alegerea şi reglajul echipamentelor s-au făcut "
          "conform proiectului; dispozitivele de separare şi comandă au fost amplasate în locurile "
          "corespunzătoare; culorile de identificare a conductoarelor au fost folosite conform "
          "normativului; conexiunile conductoarelor au fost realizate corect."),
    ("p", "Verificările prin încercări se execută, în măsura în care sunt aplicabile, de "
          "preferinţă în următoarea ordine:"),
    ("li", "continuitatea conductoarelor de protecţie şi a legăturilor echipotenţiale;"),
    ("li", "rezistenţa de izolaţie a conductoarelor şi cablurilor electrice;"),
    ("li", "separarea circuitelor;"),
    ("li", "protecţia prin deconectarea automată a alimentării (inclusiv funcţionarea "
           "dispozitivelor diferenţiale prin butonul de test);"),
    ("li", "măsurarea rezistenţei de dispersie a prizei de pământ."),
    ("p", "Punerea în funcţiune se face obligatoriu numai după efectuarea verificărilor menţionate "
          "şi întocmirea buletinelor de verificare. După punerea în funcţiune se verifică modul de "
          "funcţionare al tuturor instalaţiilor de iluminat şi prize din clădire."),
    ("p", "Recepţia lucrărilor se face conform reglementărilor în vigoare privind recepţia "
          "lucrărilor de construcţii şi instalaţii aferente. La recepţie, executantul va proba prin "
          "documente tehnice legale calitatea materialelor folosite şi execuţia corectă a "
          "lucrărilor ascunse (procese-verbale de lucrări ascunse), precum şi rezultatele probelor "
          "prevăzute a se executa înaintea, în timpul şi la terminarea lucrării. Prin recepţie, "
          "constructorul rămâne cu obligaţia remedierii eventualelor deficienţe constatate sau "
          "ivite ulterior, ca urmare a unor vicii ascunse, în perioada de garanţie."),
]

_CS_CAP6 = [
    ("h1", "6. INSTRUCŢIUNI DE TEHNICA SECURITĂŢII MUNCII"),
    ("p", "Prezentele instrucţiuni prezintă principalele măsuri de protecţia muncii care trebuie "
          "respectate la montajul, verificarea, punerea în funcţiune, exploatarea şi întreţinerea "
          "instalaţiilor electrice. Instrucţiunile nu sunt limitative; executantul va îndeplini "
          "toate obligaţiile privind securitatea şi sănătatea în muncă din actele normative în "
          "vigoare (Legea 319/2006 şi normele metodologice de aplicare)."),
    ("p", "Instruirea personalului se efectuează în fazele: instructaj de angajare, instructaj "
          "periodic şi instructaj la schimbarea locului de muncă. Obligaţia efectuării "
          "instructajului o au cei care organizează şi conduc procesul de muncă."),
    ("li", "instalaţiile electrice se execută, se verifică şi se întreţin numai de personal "
           "calificat şi autorizat; manevrele în instalaţii se execută numai de personalul de "
           "deservire operativă;"),
    ("li", "se interzice lucrul la circuitele electrice aflate sub tensiune; în punctul în care se "
           "realizează scoaterea de sub tensiune se montează indicatoare de securitate;"),
    ("li", "toate sculele şi utilajele alimentate la tensiuni peste 24 V vor avea obligatoriu "
           "conductor de legare la pământ; echipamentele se fixează definitiv în suporţi imediat "
           "după montare şi se leagă la pământ;"),
    ("li", "se interzice utilizarea conductoarelor instalaţiei de protecţie drept conductoare de "
           "fază sau de nul de lucru şi conectarea în serie la instalaţia de legare la pământ a "
           "mai multor elemente;"),
    ("li", "la înălţimi de peste 2 m, exceptând platformele stabile şi sigure, toate lucrările se "
           "execută cu centură de siguranţă; zonele cu pericol de accidentare se semnalizează cu "
           "indicatoare de avertizare;"),
    ("li", "personalul va refuza executarea lucrărilor dacă nu se asigură dotarea cu mijloacele de "
           "protecţie necesare; mijloacele de protecţie individuală se păstrează, se întreţin şi se "
           "prezintă periodic la control;"),
    ("li", "în caz de incendiu la instalaţiile electrice, înainte de a se acţiona pentru stingere, "
           "instalaţiile afectate se scot de sub tensiune; pentru stingere se folosesc numai "
           "stingătoare cu pulbere sau cu dioxid de carbon."),
    ("p", "Aceste instrucţiuni vor fi completate de conducerea tehnică a unităţilor de montaj şi "
          "exploatare şi constituie baza de instruire a personalului care lucrează în instalaţiile "
          "electrice."),
]


# =============================================================================
# Secţiuni DINAMICE
# =============================================================================

_DIAC = str.maketrans("ĂÂÎŞȘŢȚăâîşșţț", "AAISSTTaaisstt")


def _plansa_forta_ref(planse):
    """Referinţa dinamică la planşa de forţă (cap. 4): 'planşei IE.2 — PLAN ... DE FORŢĂ'.
    Lipsă -> formulare generică (fail-safe)."""
    try:
        for p in planse or []:
            titlu = str((p or {}).get("titlu") or "")
            if "FORTA" in titlu.upper().translate(_DIAC):
                nr = str(p.get("nr") or "").strip()
                if nr:
                    return "planşei {} — {}".format(nr, titlu.strip())
    except Exception:
        pass
    return "planşelor din prezentul proiect"


def _mentiune_fv(solar):
    """Menţiunea scurtă FV (cap. 1) — descriptiv, din ce e format, FĂRĂ calcule.
    Gol / eroare -> None (fără menţiune)."""
    try:
        if not isinstance(solar, dict):
            return None
        kw_raw = solar.get("package_kw") or solar.get("power_kw")
        if not kw_raw:
            return None
        from schema_fv import FV_PACKAGES, snap_fv_package
        kw = snap_fv_package(kw_raw)
        pkg = FV_PACKAGES.get(kw) or {}
        nrp, pi = pkg.get("nr_panouri"), pkg.get("pi_kw")
        if nrp and pi:
            return ("Obiectivul este prevăzut cu un sistem fotovoltaic de {} kW, compus din {} "
                    "panouri fotovoltaice (putere instalată {} kWp), invertor solar trifazat de {} kW "
                    "şi echipamentele de protecţie c.c./c.a. aferente, conform planşei şi schemei "
                    "monofilare FV din proiect. Execuţia şi racordarea sistemului fotovoltaic se vor "
                    "realiza de personal autorizat, cu respectarea capitolului dedicat din prezentul "
                    "caiet de sarcini şi a reglementărilor ANRE privind prosumatorii."
                    .format(kw, nrp, str(pi).replace(".", ","), kw))
        return ("Obiectivul este prevăzut cu un sistem fotovoltaic de {} kW, conform planşei şi "
                "schemei monofilare FV din proiect.".format(kw))
    except Exception:
        return None


def _emit_blocks(doc, blocks, alimentare=None):
    _firida = _alim_din_firida(alimentare)
    for kind, text in blocks:
        # alimentare din firidă -> fraza de branşament se înlocuieşte (ancorată pe începutul ei)
        if _firida and kind == "p" and text.startswith(_ALIM_CAIET_ANCORA):
            text = _ALIM_CAIET_TEXT
        if kind == "h1":
            _add_heading(doc, text, level=1)
        elif kind == "h2":
            _add_heading(doc, text, level=2)
        elif kind == "li":
            _add_para(doc, "- " + text)
        else:
            _add_para(doc, text)


def _nominalizari_planse(doc, planse):
    """1.4 — lista planşelor REALE din proiect. Goală -> trimitere la borderou."""
    _add_heading(doc, "1.4. Nominalizări planşe", level=2)
    rows = [p for p in (planse or []) if isinstance(p, dict) and (p.get("nr") or p.get("titlu"))]
    if not rows:
        _add_para(doc, "- conform borderoului de piese desenate al proiectului")
        return
    _add_para(doc, "Prezentul caiet de sarcini se citeşte împreună cu următoarele planşe:")
    for p in rows:
        _add_para(doc, "- {} — {}".format(str(p.get("nr") or "").strip() or "IE.-",
                                          str(p.get("titlu") or "").strip()))


_CABLU_RE = re.compile(r"(\d+)\s*[x×]\s*(\d+(?:[.,]\d+)?)")


def _fmt_cablu(s):
    """Uniformizează notaţiile de cablu la formatul TEG: 'CYY-F {n}x{sec}' + sufixele utile
    (IP44). Circuitele dedicate din formular vin în alt format decât cele derivate —
    normalizarea e DOAR la randare (decizia Dan), sursa `circuits` rămâne neatinsă.
    '3x1.5 mm2 CYYF' / 'CYY-F 5x6mmp' / '5x4 mm2 CYYF' -> 'CYY-F 3x1.5' / 'CYY-F 5x6' / 'CYY-F 5x4'.
    Nerecunoscut -> verbatim (fail-safe)."""
    raw = str(s or "").strip()
    m = _CABLU_RE.search(raw)
    if not m:
        return raw
    n, sec = m.group(1), m.group(2).replace(",", ".")
    rest = re.sub(r"mm2|mmp|mm²", " ", raw[:m.start()] + " " + raw[m.end():], flags=re.I)
    fam, sufixe = None, []
    for t in [t for t in re.split(r"[\s,;/]+", rest) if t]:
        tu = t.upper()
        if tu.startswith("IP"):
            sufixe.append(tu)                       # grad de protecţie (IP44) — sufix, nu familie
        elif fam is None and tu in ("CYYF", "CYY-F", "CYY"):
            fam = "CYY-F"
        elif fam is None and re.fullmatch(r"[A-Z][A-Z0-9-]{2,}", tu):
            fam = tu                                # altă familie (CYABY, MYF...) — păstrată
        else:
            sufixe.append(t)
    out = "{} {}x{}".format(fam or "CYY-F", n, sec)
    return (out + " " + " ".join(sufixe)).strip()


def _fmt_protectie(c):
    bt = str(c.get("breaker_type") or "").strip()
    ba = c.get("breaker_a")
    prot = " ".join(x for x in (bt, "{}A".format(ba) if ba else "") if x).strip() or "-"
    rccb = c.get("rccb_ma")
    if rccb and "RCCB" not in prot.upper():
        prot += " + RCCB {}mA".format(rccb)
    return prot


def _specificatie_tablouri(doc, circuits):
    """Cap. 7 — tabel per tablou (TEG / TES n / TE-CT) din circuitele REALE."""
    _add_heading(doc, "7. SPECIFICAŢIE TABLOURI ELECTRICE", level=1)
    _add_para(doc, "Tablourile electrice se vor echipa conform specificaţiilor de mai jos şi pe "
                   "baza schemelor electrice monofilare din proiect. Se pot folosi echipamente de "
                   "diferite provenienţe, cu condiţia să corespundă caracteristicilor tehnice "
                   "prevăzute în proiect şi să posede agrementele tehnice necesare.")
    cs = [c for c in (circuits or []) if isinstance(c, dict) and c.get("id")]
    if not cs:
        _add_para(doc, "- echiparea tablourilor: conform schemelor monofilare anexate")
        return
    panels = []
    for c in cs:
        p = str(c.get("panel") or "TEG").upper()
        if p not in panels:
            panels.append(p)
    # ordinea: TEG, TES* (natural), TE-CT
    panels.sort(key=lambda p: (0 if p == "TEG" else 2 if p.startswith("TE-CT") else 1, p))
    col_w = (Cm(2.2), Cm(6.8), Cm(4.6), Cm(3.4))
    for panel in panels:
        rows = [c for c in cs if str(c.get("panel") or "TEG").upper() == panel]
        if not rows:
            continue
        titlu = {"TEG": "Tablou electric general (TEG)",
                 "TE-CT": "Tablou electric cameră tehnică (TE-CT)"}.get(
            panel, "Tablou electric secundar ({})".format(panel))
        _add_heading(doc, titlu, level=2)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.allow_autofit = False
        _set_table_fixed(tbl)
        for col, w in zip(tbl.columns, col_w):
            col.width = w
        hdr = tbl.rows[0].cells
        for cell, w, txt in zip(hdr, col_w, ("Circuit", "Destinaţie", "Protecţie", "Cablu")):
            cell.width = w
            _set_run_font(cell.paragraphs[0].add_run(txt), size=10, bold=True)
        for c in rows:
            cells = tbl.add_row().cells
            vals = (str(c.get("id") or ""),
                    str(c.get("description") or c.get("usage") or ""),
                    _fmt_protectie(c),
                    _fmt_cablu(c.get("cable_type") or c.get("cable")))
            for cell, w, txt in zip(cells, col_w, vals):
                cell.width = w
                _set_run_font(cell.paragraphs[0].add_run(txt), size=10, bold=False)
        set_table_borders(tbl)
        doc.add_paragraph()


# CURENTI SLABI — cerinte de EXECUTIE. Spre deosebire de iluminatul de siguranta (text static,
# fiindca reguli de montaj exista oricum), aici textul apare doar cand proiectul chiar are
# echipamente: cerinta de non-regresie e ca proiectele fara curenti slabi sa iasa neschimbate.
# Inaltimile sunt cele din proiect (aceleasi valori ca `_CS_HEIGHT` din draw_elements).
_CS_DATE_TV = ("priza_date", "priza_tv", "priza_mixta")
_CS_ACTIVE = ("centrala_efractie", "tastatura_efractie", "detector_pir", "contact_magnetic",
              "sirena_interioara", "sirena_exterioara", "buton_panica", "camera_video", "nvr")


def _topologie_stea_cerinte(comp):
    """Cerinţele de topologie: stea, panou de conexiuni, limita de 90 m pe legătura orizontală."""
    try:
        from draw_elements import cs_utp_cabluri, cs_patch_bom
    except Exception:
        return ""
    n = cs_utp_cabluri(comp)
    if not n:
        return ""
    _pp = cs_patch_bom(n)
    _p = ", ".join("%d panouri de %d porturi" % (b, p) if b > 1 else "un panou de %d porturi" % p
                   for p, b in _pp)
    return ("Cablarea se execută în topologie de stea, conform SR EN 50173: de la fiecare priză de "
            "date şi de la fiecare cameră se trage un cablu propriu, continuu, până la panoul de "
            "conexiuni din rack — fără înnădiri şi fără derivaţii pe traseu. Se prevăd %d legături "
            "şi %s, cu porturile numerotate şi etichetate corespunzător prizei deservite. Lungimea "
            "unei legături orizontale nu depăşeşte 90 m; peste această lungime soluţia se "
            "reproiectează cu un punct de distribuţie intermediar. Legătura dintre panoul de "
            "conexiuni şi echipamentul activ se face cu cordoane de echipare (patch cord) de "
            "aceeaşi categorie cu cablul instalat." % (n, _p))


def _video_alimentare_cerinte(comp):
    """Cerinţele de alimentare a camerelor. PoE-ul e soluţia de BAZĂ (decizia Dan): înregistratorul
    se specifică cu PoE integrat, dimensionat pe camerele plasate. Cablul separat de 2×1 mmp la
    camere nu mai este necesar; sursa cu acumulator rămâne a instalaţiei de efracţie."""
    try:
        from draw_elements import cs_nvr
    except Exception:                      # fail-safe: cerinţa generică, fără dimensionare
        return ("Alimentarea camerelor se face prin PoE, din înregistratorul montat în rack, pe "
                "acelaşi cablu UTP pe care transmit imaginea.")
    n_cam = int((comp or {}).get("camera_video") or 0)
    t = cs_nvr(n_cam)
    if not n_cam:
        return ("Alimentarea camerelor se face prin PoE, din înregistratorul montat în rack, pe "
                "acelaşi cablu UTP pe care transmit imaginea.")
    if not t["poe"]:
        return ("Înregistratorul se prevede cu %d canale, pentru cele %d camere din proiect. La "
                "această mărime înregistratoarele nu dispun de alimentare PoE integrată, astfel "
                "încât alimentarea camerelor se face prin comutatoare PoE montate în rack, pe "
                "acelaşi cablu UTP pe care camerele transmit imaginea. Bugetul de putere al "
                "comutatoarelor se verifică faţă de consumul însumat al camerelor alimentate, cu "
                "o rezervă de minimum 20%%. Camerele PTZ se alimentează pe porturi PoE+ (IEEE "
                "802.3at), care asigură 30 W pe port." % (t["canale"], n_cam))
    _t = ("Înregistratorul se prevede cu %d canale şi cu alimentare PoE integrată, pentru cele %d "
          "camere din proiect. Alimentarea camerelor se face prin PoE, pe acelaşi cablu UTP pe care "
          "transmit imaginea, direct din porturile înregistratorului; nu se prevăd circuite "
          "separate de alimentare la camere. Camerele PTZ se alimentează pe porturi PoE+ (IEEE "
          "802.3at, 30 W pe port), celelalte pe PoE standard (IEEE 802.3af, 15,4 W pe port). "
          "Bugetul PoE al înregistratorului se verifică faţă de consumul însumat al camerelor "
          "alimentate, cu o rezervă de minimum 20%%." % (t["canale"], n_cam))
    _fara = max(0, n_cam - t["poe"])
    if _fara:
        _t += (" Înregistratorul dispune de %d porturi PoE; cele %d camere rămase se alimentează "
               "dintr-un comutator PoE suplimentar, montat în rack şi alimentat din acelaşi "
               "circuit de 230 V." % (t["poe"], _fara))
    _t += (" Sursa cu acumulator din rack rămâne destinată instalaţiei de efracţie şi se "
           "dimensionează numai pe consumul acesteia.")
    return _t


def _dtv_echipament_cerinte(comp):
    """Cerinţele pentru echipamentul de distribuţie, DIMENSIONAT din numărul de prize plasate.
    Dimensionarea vine din `draw_elements` — aceeaşi sursă ca lista de cantităţi, memoriul şi schema."""
    try:
        from draw_elements import cs_prize_dtv, cs_switch_porturi, cs_splitter_bom
    except Exception:
        return ""                          # fail-safe: cerinţa lipseşte, restul blocului rămâne
    n_date, n_tv = cs_prize_dtv(comp)
    porturi, spl = cs_switch_porturi(n_date), cs_splitter_bom(n_tv)
    if not porturi and not spl:
        return ""
    _t = "În rack se montează "
    _b = []
    if porturi:
        _b.append("switch-ul de date cu %d porturi, dimensionat pentru cele %d prize de date "
                  "deservite plus portul de legătură către router" % (porturi, n_date))
    if spl:
        if sum(n for _, n in spl) == 1:
            _b.append("splitterul TV pasiv cu %d ieşiri, pentru cele %d prize de televiziune"
                      % (spl[0][0], n_tv))
        else:
            _b.append("splitterele TV pasive montate în cascadă (%s), pentru cele %d prize de "
                      "televiziune" % (", ".join("%d bucăţi cu %d ieşiri" % (n, i) if n > 1
                                                 else "unul cu %d ieşiri" % i for i, n in spl), n_tv))
    _t += " şi ".join(_b) + ". "
    if porturi:
        _t += ("Switch-ul se alimentează din circuitul dedicat de 230 V al punctului de "
               "distribuţie. ")
    if spl:
        # cerinta reala de montaj la pasive: iesirile libere nemarcate dezadapteaza linia si intorc
        # semnal in retea (unde reflexiile se vad ca „fantome" pe imagine)
        _t += ("Splitterele sunt pasive şi nu se alimentează; la montaj se respectă sensul "
               "intrare/ieşiri marcat pe carcasă, iar ieşirile neutilizate se închid cu sarcină de "
               "75 ohm. ")
    _t += ("Echipamentele se montează în rack cu acces frontal la conectori şi se etichetează, "
           "fiecare port fiind marcat cu identificatorul prizei pe care o deserveşte.")
    return _t


def _cerinte_curenti_slabi(doc, comp):
    """Blocul de executie din cap. 3. Nu scrie nimic fara echipamente pe plan.
    Cerintele de montaj/PIF privesc echipamentele ACTIVE (efractie + video); la o casa cu doar prize
    de date si TV se scrie doar ce se aplica lor — nu inaltimi de sirene si nu programare de
    centrala, care n-ar avea obiect."""
    if not comp:
        return
    _activ = any(comp.get(k) for k in _CS_ACTIVE)
    _dtv = any(comp.get(k) for k in _CS_DATE_TV)
    _add_heading(doc, "Instalaţii de curenţi slabi", level=2)
    _top = _topologie_stea_cerinte(comp)
    if _top:
        _add_para(doc, _top)
    if _dtv:
        _add_para(doc, "Prizele de date şi de televiziune se montează în doze proprii, la "
                       "înălţimile din planşele proiectului: prizele de date la 0,30 m faţă de "
                       "pardoseală, alături de prizele de 230 V, iar cele de televiziune la 1,20 m, "
                       "în spatele televizorului. Legăturile se fac radial, câte un cablu de la "
                       "fiecare priză până la punctul de distribuţie — fără înnădiri şi fără "
                       "derivaţii pe traseu. Cablul de date este UTP cat. 5e/6, terminat pe "
                       "conectori RJ45 după schema T568B la ambele capete; cel de televiziune este "
                       "coaxial de 75 ohm, cu conectori F. La recepţie se verifică fiecare priză "
                       "prin testare de continuitate şi de perechi, iar prizele TV prin măsurarea "
                       "nivelului de semnal.")
        _e = _dtv_echipament_cerinte(comp)
        if _e:
            _add_para(doc, _e)
    if not _activ:
        return
    
    _add_para(doc, "Echipamentele de curenţi slabi (detecţie efracţie şi supraveghere video) se "
                   "montează la înălţimile prevăzute în planşele proiectului: detectoarele de "
                   "mişcare şi camerele de supraveghere la 2,5 m faţă de pardoseală, centrala de "
                   "efracţie la 1,5 m, tastatura de comandă la 1,4 m, butonul de panică la 1,3 m, "
                   "iar sirena exterioară la 3,0 m, pe faţadă, în poziţie vizibilă şi greu "
                   "accesibilă. Poziţiile din planşe se respectă; orice modificare impusă de "
                   "situaţia din teren se face numai cu acordul scris al proiectantului.")
    _add_para(doc, "Circuitele de curenţi slabi se pozează separat de cele de curenţi tari, în "
                   "tuburi de protecţie proprii sau pe trasee distincte, respectând distanţa minimă "
                   "faţă de conductoarele de 230 V prevăzută în tabelul 3.1, art. 3.5 din "
                   "Normativul I7-2011. Este interzisă tragerea cablurilor de semnal în acelaşi tub "
                   "cu conductoarele de forţă sau de iluminat. Traversările inevitabile ale "
                   "traseelor de curenţi tari se execută în unghi drept, cu păstrarea distanţei de "
                   "izolare.")
    # Cele trei tipuri sunt EXACT cele din `_CS_CABLE` (planşa + lista de cantitati) si descriu
    # ACELASI sistem ca memoriul: camere IP -> UTP catre NVR. Coaxialul RG59 (CCTV analogic) a fost
    # scos din toate locurile la corectia din 30 aug 2026.
    _add_para(doc, "Cablurile utilizate sunt: UTP cat. 5e/6 pentru transmisia de date şi pentru "
                   "alimentarea PoE a camerelor video IP către înregistrator, cablu de alimentare "
                   "2×1 mmp pentru echipamentele de efracţie alimentate în 12 V c.c. şi cablu de "
                   "semnal ecranat 2×(LiY(St)Y) 3×2×0,6 pentru buclele de "
                   "detecţie. Cablurile se pozează fără "
                   "înnădiri pe traseu; legăturile se fac exclusiv în dozele prevăzute în proiect, "
                   "cu cleme, şi rămân accesibile. Ecranele cablurilor de semnal se leagă la masă "
                   "într-un singur punct, la centrală, pentru a nu forma bucle de masă.")
    # PoE INVERSAT (decizia Dan): era alternativa, devine solutia de BAZA. Inregistratorul se
    # specifica cu PoE integrat si dimensionat pe camere, deci cablul separat de 2x1 mmp la camere
    # dispare. Sursa cu acumulator RAMANE, dar a efractiei — de-aia si randul ei de legenda s-a
    # schimbat de la „instalatie monitorizare video, max. 10 camere".
    _add_para(doc, _video_alimentare_cerinte(comp))
    _add_para(doc, "După montaj, executantul realizează punerea în funcţiune şi programarea "
                   "centralei de efracţie: definirea zonelor şi a partiţiilor, temporizările de "
                   "intrare şi ieşire, codurile de utilizator şi codul de instalator, precum şi "
                   "configurarea înregistratorului video — rezoluţie, cadenţă, durata de stocare şi "
                   "declanşarea la mişcare. Parolele implicite din fabrică se schimbă obligatoriu "
                   "înainte de predare.")


def _receptie_curenti_slabi(doc, comp):
    """Completarea la cap. 5 — ce se verifica la receptie pe partea de curenti slabi.
    Verificarile de mai jos privesc echipamentele ACTIVE; verificarea prizelor de date/TV (test de
    perechi, nivel de semnal) e in blocul de executie, langa cerintele lor."""
    if not comp or not any(comp.get(k) for k in _CS_ACTIVE):
        return
    _add_para(doc, "La recepţia instalaţiilor de curenţi slabi se verifică suplimentar: "
                   "funcţionarea fiecărui detector şi a fiecărui contact magnetic, prin declanşare "
                   "individuală şi confirmarea semnalizării la centrală; funcţionarea sirenelor "
                   "interioară şi exterioară; funcţionarea butoanelor de panică; autonomia "
                   "acumulatorului de rezervă, prin proba de întrerupere a alimentării de 230 V şi "
                   "menţinerea sistemului în funcţiune pe durata declarată; calitatea imaginii "
                   "fiecărei camere, ziua şi noaptea, cu verificarea câmpului de acoperire faţă de "
                   "cel figurat în planşă; înregistrarea şi redarea corectă pe NVR.")
    _add_para(doc, "La predare, executantul înmânează beneficiarului documentaţia sistemului: "
                   "codurile de utilizator şi de instalator, manualele de utilizare ale centralei "
                   "şi ale înregistratorului video, schema zonelor programate, certificatele de "
                   "garanţie ale echipamentelor şi procesul-verbal de punere în funcţiune. "
                   "Beneficiarul este instruit asupra armării şi dezarmării sistemului şi asupra "
                   "procedurii în caz de alarmă falsă.")


# ALIMENTARE DIN FIRIDĂ: la un spaţiu dintr-un imobil existent nu există branşament al obiectivului
# şi nici proiect al furnizorului — racordul din firidă ţine de administratorul imobilului. Fraza se
# înlocuieşte, ancorată pe începutul ei. Fără câmpul `alimentare` -> caietul iese identic cu HEAD.
_ALIM_CAIET_ANCORA = "Instalaţiile electrice de utilizare se vor executa numai de către firme"
_ALIM_CAIET_TEXT = (
    "Instalaţiile electrice de utilizare se vor executa numai de către firme atestate şi/sau "
    "electricieni autorizaţi conform ordinelor ANRE, având gradul de competenţă corespunzător "
    "lucrării executate. Instalaţiile electrice se vor executa cu respectarea normelor şi "
    "reglementărilor în vigoare. Spaţiul se alimentează din firida de distribuţie a imobilului, "
    "printr-un racord care ţine de instalaţia existentă a acestuia; execuţia lui se face cu acordul "
    "administratorului imobilului, iar puterea instalată rezultată din proiect se declară acestuia "
    "şi furnizorului de energie, pentru verificarea disponibilităţii pe coloana existentă."
)


def _pagina_finala(doc, cf):
    """Ultima pagină: PROIECTANT + FIRMĂ, jos-dreapta (datele din cartuşul proiectului)."""
    doc.add_page_break()
    for _ in range(18):                     # împinge blocul spre partea de jos a paginii
        doc.add_paragraph()
    rows = [
        ("PROIECTANT DE SPECIALITATE", str(cf.get("firma_nume") or "")),
        ("ŞEF PROIECT", str(cf.get("sef_proiect") or "")),
        ("ÎNTOCMIT", str(cf.get("proiectant_nume") or "")),
    ]
    tw_l, tw_r = Cm(5.2), Cm(4.6)
    tbl = doc.add_table(rows=0, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.RIGHT
    tbl.allow_autofit = False
    _set_table_fixed(tbl)
    for col, w in zip(tbl.columns, (tw_l, tw_r)):
        col.width = w
    for label, value in rows:
        cells = tbl.add_row().cells
        cells[0].width, cells[1].width = tw_l, tw_r
        _set_run_font(cells[0].paragraphs[0].add_run(label), size=11, bold=True)
        _set_run_font(cells[1].paragraphs[0].add_run(value), size=11, bold=False)
    set_table_borders(tbl)


# =============================================================================
# Build
# =============================================================================

def build_caiet_docx(data: dict) -> bytes:
    """Construieşte caietul de sarcini .docx şi returnează bytes.
    Gate-ul de fază (doar DTAC+PT / PT) e responsabilitatea endpoint-ului."""
    data = data or {}
    cp = data.get("cartus_proiect") or {}
    cf = data.get("cartus_firma") or {}
    planse = data.get("planse") or []
    circuits = data.get("circuits") or []
    solar = data.get("solar") or {}
    # CURENŢI SLABI: gate pe PREZENŢĂ REALĂ, din circuite (`_cs_componente`, pus de enrich_circuits).
    # Fără echipamente pe plan -> caietul iese byte-identic ca înainte.
    _cs_comp, _ = _cs_din_circuite(circuits)
    # ALIMENTAREA: "din_firida" -> fraza de racord din firidă; altfel (inclusiv absent) -> ca azi.
    alimentare = data.get("alimentare") or ""

    # Lista REALĂ de planşe (aceeaşi autoritate ca borderoul memoriului la PT):
    # planşele explicite din payload; altfel derivate din numerotare (fail-safe: cele primite).
    try:
        from plansa_numbering import compute_plansa_numbering, derive_extra_floors
        has_tect = data.get("has_tect")
        if has_tect is None:
            has_tect = any((c or {}).get("panel") == "TE-CT" for c in circuits)
        extra = data.get("extra_floors")
        if extra is None:
            extra = derive_extra_floors(circuits)
        has_fv = bool(solar.get("package_kw") or solar.get("power_kw"))
        real = compute_plansa_numbering(extra, bool(has_tect), has_fv=has_fv)
        if real:
            planse = [{"nr": p["nr"], "titlu": p["nume"]} for p in real]
    except Exception:
        pass

    doc = _setup_document()
    _page_coperta(doc, cp, cf, titlu="CAIET DE SARCINI INSTALAȚII ELECTRICE")

    _add_heading(doc, "CAIET DE SARCINI INSTALAŢII ELECTRICE", level=1)

    _emit_blocks(doc, _CS_CAP1, alimentare)
    _nominalizari_planse(doc, planse)               # 1.4 dinamic
    _emit_blocks(doc, _CS_SARCINI)                  # 1.5 + 1.6
    fv_txt = _mentiune_fv(solar)
    if fv_txt:
        _add_heading(doc, "Sistem fotovoltaic", level=2)
        _add_para(doc, fv_txt)

    _emit_blocks(doc, _CS_NORMATIVE)                # 2 (lista validată)
    if fv_txt:
        _emit_blocks(doc, _CS_NORMATIVE_FV)         # normele FV, doar cu FV

    _emit_blocks(doc, _CS_CAP3)                     # 3
    _cerinte_curenti_slabi(doc, _cs_comp)           # 3.x, doar cu echipamente pe plan

    _add_heading(doc, "4. EXECUTAREA INSTALAŢIILOR DE LEGARE LA PĂMÂNT", level=1)
    # Formularea EXACTĂ a lui Dan + referinţa dinamică la planşă.
    _add_para(doc, "Instalaţia de legare la pământ este obligatorie conform I7-2011 şi se execută "
                   "cu platbandă OL-Zn 40×4 mm, conform " + _plansa_forta_ref(planse) + ".")
    _emit_blocks(doc, _CS_CAP4_PRINCIPII)

    _emit_blocks(doc, _CS_CAP5)                     # 5
    _receptie_curenti_slabi(doc, _cs_comp)          # completarea de receptie, doar cu echipamente
    _emit_blocks(doc, _CS_CAP6)                     # 6
    _specificatie_tablouri(doc, circuits)           # 7 dinamic

    _pagina_finala(doc, cf)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
