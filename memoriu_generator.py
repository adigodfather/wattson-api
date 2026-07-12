# -*- coding: utf-8 -*-
"""
Generator memoriu tehnic instalatii electrice (.docx) pentru ZYNAPSE.

Construieste un document Word completat automat din datele cartusului de proiect.
Schelet functional: coperta + fisa proiectului + borderou + schelet memoriu.
Textul fix lung al sectiunilor tehnice este momentan un PLACEHOLDER
("[TEXT_FIX_MEMORIU_AICI]") — va fi inlocuit ulterior.

Returneaza bytes (.docx) — encodarea base64 se face in endpoint.
"""

import io

import requests
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "Arial"

# Latimi tabel info (2 coloane) — total ~16cm, centrat pe pagina
COL_LEFT_W = Cm(5.5)
COL_RIGHT_W = Cm(10.5)
CELL_BG_GREY = "D9D9D9"


# =============================================================================
# HELPERS — font, shading, tabele, paragrafe
# =============================================================================

def _set_run_font(run, size=11, bold=False):
    """Forteaza Arial (ascii/hAnsi/cs) + dimensiune pe un run."""
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rfonts.set(qn("w:cs"), FONT_NAME)
    return run


def _set_style_font(style, size=None, bold=None, color=None):
    """Aplica Arial (+ optional size/bold/color) pe un stil de document."""
    style.font.name = FONT_NAME
    if size is not None:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rfonts.set(qn("w:cs"), FONT_NAME)


def set_cell_bg(cell, color_hex):
    """Fundal celula (shading XML w:shd)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _set_table_fixed(table):
    """Layout fix (python-docx e capricios la latimi fara asta)."""
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)


def set_table_borders(table):
    """Chenare negre vizibile pe toate laturile (tblBorders XML)."""
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement("w:{}".format(edge))
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")        # ~1pt
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tbl_pr.append(borders)


def _blank(doc, n=1):
    for _ in range(n):
        doc.add_paragraph()


def _add_centered(doc, text, size=11, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(p.add_run(text), size=size, bold=bold)
    return p


def _add_heading(doc, text, level=1, centered=False):
    """Heading pe stilul nativ (deja setat Arial) — cu override pe run pentru siguranta."""
    style = "Heading 1" if level == 1 else "Heading 2"
    p = doc.add_paragraph(style=style)
    if centered:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    size = 14 if level == 1 else 12
    _set_run_font(p.add_run(text), size=size, bold=True)
    return p


def _add_label_value(doc, label, value="", size=11):
    """Paragraf cu eticheta bold + valoare normala."""
    p = doc.add_paragraph()
    _set_run_font(p.add_run(label), size=size, bold=True)
    if value:
        _set_run_font(p.add_run(value), size=size, bold=False)
    return p


def _add_para(doc, text, size=11, bold=False, align=None):
    """Paragraf normal (optional aliniere)."""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    _set_run_font(p.add_run(text), size=size, bold=bold)
    return p


def _add_info_table(doc, rows, bg=CELL_BG_GREY):
    """Tabel 2 coloane centrat: stanga gri+bold (eticheta), dreapta bold (valoare).
    Seteaza atat latimea coloanelor cat si a fiecarei celule (python-docx capricios).
    """
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    _set_table_fixed(table)

    # latimi pe coloane
    for col, w in zip(table.columns, (COL_LEFT_W, COL_RIGHT_W)):
        col.width = w

    for label, value in rows:
        cells = table.add_row().cells
        c_left, c_right = cells[0], cells[1]
        c_left.width = COL_LEFT_W
        c_right.width = COL_RIGHT_W
        set_cell_bg(c_left, bg)

        p_left = c_left.paragraphs[0]
        _set_run_font(p_left.add_run(label), size=11, bold=True)

        p_right = c_right.paragraphs[0]
        _set_run_font(p_right.add_run(value), size=11, bold=True)

    set_table_borders(table)  # chenare vizibile (coperta + fisa proiectului)
    return table


def _try_add_logo(doc, logo_url):
    """Insereaza logo centrat (~5cm) daca url e nevid si download-ul reuseste.
    Esecul nu propaga eroare (logo optional)."""
    if not logo_url:
        return
    try:
        resp = requests.get(logo_url, timeout=10)
        resp.raise_for_status()
        stream = io.BytesIO(resp.content)
        doc.add_picture(stream, width=Cm(5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        # logo optional — omite fara eroare
        pass


# =============================================================================
# SETUP DOCUMENT — font implicit, pagina A4, margini, stiluri heading
# =============================================================================

def _setup_document():
    doc = Document()

    # Font implicit Arial 11 (stilul Normal)
    _set_style_font(doc.styles["Normal"], size=11, bold=False,
                    color=RGBColor(0, 0, 0))

    # Heading 1: Arial 14 bold negru | Heading 2: Arial 12 bold negru
    _set_style_font(doc.styles["Heading 1"], size=14, bold=True,
                    color=RGBColor(0, 0, 0))
    _set_style_font(doc.styles["Heading 2"], size=12, bold=True,
                    color=RGBColor(0, 0, 0))

    # A4 portrait + margini (sus/jos/dreapta 2cm, stanga 2.5cm)
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.left_margin = Cm(2.5)

    return doc


# =============================================================================
# PAGINI
# =============================================================================

def _page_coperta(doc, cp, cf):
    # 1. spatiere de sus
    _blank(doc, 4)

    # 2. logo (optional)
    _try_add_logo(doc, cf.get("firma_logo_url", ""))

    # 3-5. identitate firma
    _add_centered(doc, cf.get("firma_nume", ""), size=14, bold=True)
    _add_centered(
        doc,
        "Reg. Com. {}   ·   C.U.I. {}".format(
            cf.get("firma_reg_com", ""), cf.get("firma_cui", "")),
        size=9,
    )
    _add_centered(
        doc,
        "Tel. {}   ·   {}".format(
            cf.get("firma_tel", ""), cf.get("firma_email", "")),
        size=9,
    )

    # 6. spatiu mare
    _blank(doc, 6)

    # 7. titlu memoriu
    _add_centered(doc, "MEMORIU TEHNIC INSTALAȚII ELECTRICE",
                  size=16, bold=True)

    # 8. spatiu
    _blank(doc, 3)

    # 9. tabel date proiect
    _add_info_table(doc, [
        ("BENEFICIAR:", cp.get("beneficiar", "")),
        ("LUCRARE:", cp.get("titlu_proiect", "")),
        ("ADRESA:", cp.get("amplasament", "")),
        ("PROIECT:", "Nr. {}".format(cp.get("numar_proiect", ""))),
        ("FAZA:", cp.get("faza", "")),
    ])

    # 10. page break
    doc.add_page_break()


def _page_fisa(doc, cp, cf):
    # impins spre mijloc-sus
    _blank(doc, 4)
    _add_heading(doc, "I. FIȘA PROIECTULUI", level=1, centered=True)
    _blank(doc, 1)

    proiectant_val = (
        "{firma_nume}  Nr. Reg. ONRC: {reg}  CUI: {cui}  Tel.: {tel}  "
        "e-mail: {email}  {atestat}".format(
            firma_nume=cf.get("firma_nume", ""),
            reg=cf.get("firma_reg_com", ""),
            cui=cf.get("firma_cui", ""),
            tel=cf.get("firma_tel", ""),
            email=cf.get("firma_email", ""),
            atestat=cf.get("firma_atestat", ""),
        )
    )

    _add_info_table(doc, [
        ("FAZA DE PROIECTARE:", cp.get("faza", "")),
        ("LUCRARE:", cp.get("titlu_proiect", "")),
        ("AMPLASAMENT:", cp.get("amplasament", "")),
        ("BENEFICIAR:", cp.get("beneficiar", "")),
        ("VOLUM/OBIECT:", "INSTALAȚII ELECTRICE"),
        ("PROIECTANT DE SPECIALITATE:", proiectant_val),
    ])

    doc.add_page_break()


def _is_pt(faza) -> bool:
    """True dacă faza include PT (DTAC+PT / PT / 'D.T.A.C + P.T.'). DTAC pur -> False.
    Gate pentru secțiunile noi de PT (brevier, cerințe, faze determinante, program control)."""
    s = "".join(c for c in str(faza or "").upper() if c.isalpha())
    return "PT" in s


# Piese desenate STANDARD pentru faza PT (din exemplul de memoriu PT — IE.1..IE.8).
_PIESE_DESENATE_PT = [
    ("IE.1", "PLAN PARTER INSTALAȚII ELECTRICE DE ILUMINAT"),
    ("IE.2", "PLAN PARTER INSTALAȚII ELECTRICE DE FORȚĂ"),
    ("IE.3", "SCHEMA ELECTRICĂ MONOFILARĂ TABLOU ELECTRIC GENERAL"),
    ("IE.4", "SCHEMA ELECTRICĂ MONOFILARĂ CENTRALĂ TERMICĂ"),
    ("IE.5", "SCHEMA ELECTRICĂ MONOFILARĂ SISTEM FOTOVOLTAIC"),
    ("IE.6", "PLAN ÎNVELITOARE SISTEM FOTOVOLTAIC"),
    ("IE.7", "DETALIU ILUMINAT DE SIGURANȚĂ"),
    ("IE.8", "DETALIU CONECTARE PRIZĂ DE PĂMÂNT"),
]


def _page_borderou(doc, planse, is_pt=False):
    _add_heading(doc, "II. BORDEROU", level=1, centered=False)

    _set_run_font(doc.add_paragraph().add_run("PIESE SCRISE:"),
                  size=11, bold=True)
    piese_scrise = [
        "I. FIȘA PROIECTULUI",
        "II. BORDEROU",
        "III. MEMORIU TEHNIC INSTALAȚII ELECTRICE",
        "    2.1. Alimentarea cu energie electrică",
        "    2.5. Distribuția energiei electrice",
        "    2.6. Instalația de prize și forță",
        "IV. CERINȚE DE CALITATE ȘI CRITERII DE PERFORMANȚĂ",
    ]
    if is_pt:
        # secțiunile NOI de PT (conținutul lor se generează în M2-M5)
        piese_scrise += [
            "V. BREVIAR DE CALCUL",
            "VI. FAZE DETERMINANTE PENTRU INSTALAȚII ELECTRICE",
            "VII. PROGRAM DE CONTROL AL CALITĂȚII LUCRĂRILOR DE INSTALAȚII ELECTRICE",
        ]
    for line in piese_scrise:
        _set_run_font(doc.add_paragraph().add_run(line), size=11, bold=False)

    _blank(doc, 1)
    _set_run_font(doc.add_paragraph().add_run("PIESE DESENATE:"),
                  size=11, bold=True)
    # Lista REALĂ de planșe primită (nr + titlu): planuri per nivel + scheme monofilare care EXISTĂ,
    # numerotate IE.1..IE.N secvențial (calculate în build_memoriu_docx via plansa_numbering).
    # FALLBACK: dacă PT și lista lipsește -> vechea listă standard IE.1..IE.8 (backward-compat).
    piese_desenate = [(pl.get("nr", ""), pl.get("titlu", "")) for pl in (planse or [])]
    if is_pt and not piese_desenate:
        piese_desenate = _PIESE_DESENATE_PT
    for nr, titlu in piese_desenate:
        p = doc.add_paragraph()
        _set_run_font(p.add_run("{}  ".format(nr)), size=11, bold=True)
        _set_run_font(p.add_run(titlu), size=11, bold=False)

    doc.add_page_break()


# =============================================================================
# TEXT FIX MEMORIU — sectiuni tehnice (verbatim, diacritice pastrate exact)
# Tuple: (kind, text); kind in {"h1","h2","p","li"}
# =============================================================================
_MEMORIU_BLOCKS = [
    ("h2", "SITUAȚIA PROPUSĂ"),
    ("p", "Înaintea începerii lucrărilor se va obţine, prin grija beneficiarului, avizul tehnic de racordare la reţeaua furnizorului, aviz care condiţionează începerea lucrărilor de instalaţii electrice."),
    ("p", "Soluţia de branşare şi amplasarea echipamentului de măsurare a energiei electrice se va realiza în baza unui proiect tehnic elaborat conform fişei de soluţie emisă de S.D.E.E. competentă, comandat de beneficiarul lucrării."),
    ("p", "Clădirea se va alimenta cu energie electrică de la un bloc de măsură şi protecţie trifazic (BMPT) propus, amplasat la limita de proprietate. De aici se va alimenta tabloul TEG prin cablu CYABY-F 5x6 mmp. Cablul de alimentare a acestuia se va poza îngropat sub tencuială, protejat în tuburi IPEY."),
    ("p", "Se propune un tablou electric general montat la intrarea în locuinţă, din care se vor alimenta circuite de iluminat normal LED şi circuite de prize. Referitor la instalaţia de iluminat, corpurile de iluminat trebuie să fie în mod obligatoriu cu sursă LED pentru reducerea consumului de energie."),

    ("h2", "2. SOLUȚII TEHNICE"),

    ("h2", "2.1. Alimentarea cu energie electrică"),
    ("p", "Pentru diminuarea riscului de incendiu, conform art. 4.2.2.8 din I7-2011 în BMPT se va monta un dispozitiv de protecţie la curent diferenţial rezidual (DDR) cu curentul nominal de funcţionare de 300 mA. Toţi consumatorii sunt alimentaţi la tensiunea 400/230V, 50Hz."),
    ("p", "Se propune dotarea obiectivului cu un branşament trifazic, prin intermediul unui BMPT. Alimentarea cu energie electrică se face de la un bloc de măsură protecţie BMPT propus. De aici se alimentează tabloul electric general TEG, printr-un cablu CYABY-F 5x6 mmp."),
    ("p", "Din tabloul electric general se vor alimenta următoarele:"),
    ("li", "Circuite de iluminat normal"),
    ("li", "Circuite de prize"),
    ("li", "Circuit alimentare cuptor"),
    ("li", "Circuite curenţi slabi"),
    ("li", "Circuit alimentare distribuitor"),
    ("li", "Tablou electric cameră tehnică"),
    ("p", "Puterile instalate, absorbite şi curenţii absorbiţi pentru fiecare tablou sunt prezentate în schemele monofilare anexate (planşele IE)."),

    ("h2", "2.3. Priza de pământ"),
    ("p", "La prezenta clădire priza de pământ se realizează odată cu fundaţia, cu platbandă de OL-ZN 40x4 mm, dispusă pe conturul fundaţiei, înglobată direct în betonul fundaţiei clădirii, astfel încât să fie învelită cu un strat de beton de cel puţin 3 cm. Asigurarea continuităţii electrice pentru legături se face prin îmbinări sudate de bună calitate."),

    ("h2", "2.4. Instalaţiile de protecţie împotriva trăsnetului"),
    ("p", "Conform I7/2011 Cap.6, Punctul 6.2.2.6, pentru construcţia studiată nu este nevoie de instalaţie IPT. În cazul în care beneficiarul doreşte să monteze instalaţie IPT, ca măsură compensatorie, se va interveni asupra sa prin intermediul unui alt proiect."),

    ("h2", "2.5. Distribuţia energiei electrice"),
    ("p", "Pentru realizarea instalaţiei electrice la consumatori se utilizează o schemă de distribuţie combinată trifazată / monofazată cu 5 respectiv 3 conductoare. Circuitele sunt protejate la suprasarcină şi scurtcircuit prin întreruptoare automate cu declanşatoare magneto-termice şi împotriva curenţilor de defect prin dispozitive diferenţiale."),
    ("p", "Tuburile de protecţie se amplasează faţă de conductele altor instalaţii şi faţă de elementele de construcţie, respectându-se distanţele minime I7-2011. La contactul cu materiale combustibile conductoarele electrice se vor poza în tuburi sau plinte metalice sau din materiale plastice omologate pentru montare pe materiale combustibile."),

    ("h2", "2.6. Instalaţia de prize și forţă"),
    ("p", "Toate circuitele de prize trebuie să aibă protecţie diferenţială de mare sensibilitate, 30mA, DDR, pentru a asigura o protecţie suplimentară la curenţii de defect. Caracteristicile aparaturii de protecţie de pe coloanele respective sunt cuprinse în schema monofilară generală."),
    ("p", "Distribuţia circuitelor propuse se realizează cu cabluri tip CYY-F pozate aparent în paturi de cabluri şi/sau îngropat în tuburi de protecţie din IPEY-18-ST pentru circuitele de forţă, respectiv IPEY-16-ST pentru circuitele de iluminat."),
    ("p", "Conform art. 5.4.29 din Normativul I7-2011, prizele din încăperile în care au acces copii vor fi de tip special (cu obturatori) şi prevăzute cu dispozitive de protecţie diferenţială ≤ 30 mA."),
    ("p", "De menţionat că tablourile electrice vor fi de tip modular, prevăzute cu unul sau mai multe rânduri de module, fixate pe şine DIN 35 mm şi vor fi comandate de către beneficiar, pentru execuţie, testare şi montare unei firme de specialitate, pe baza documentaţiei din proiect."),

    ("h2", "2.7. Instalaţii de iluminat"),
    ("p", "Se vor realiza următoarele nivele de iluminare în stare normală, conform NP-061:"),
    ("li", "Camere: 500 lx;"),
    ("li", "Depozite, holuri, zone de circulaţie, coridoare: 100-200 lx;"),
    ("li", "Grupuri sanitare, toalete: 200 lx."),
    ("p", "Toate corpurile de iluminat vor fi cu sursă LED pentru reducerea consumului de energie. Tipul şi puterea corpurilor de iluminat se vor alege conform destinaţiei fiecărei încăperi şi nivelelor de iluminare normate."),

    ("h2", "2.8. PROTECȚIA ÎMPOTRIVA ȘOCURILOR ELECTRICE"),
    ("p", "S-au aplicat măsuri pentru protecţia utilizatorilor împotriva şocurilor electrice datorate atingerilor directe şi indirecte."),
    ("p", "Protecţia împotriva atingerilor directe se asigură prin utilizarea echipamentelor corespunzătoare categoriei de influenţe externe, conductoare izolate, tuburi de protecţie electroizolante, carcase, tablouri de distribuţie cu părţi active izolate. Se vor realiza legături de echipotenţializare cf. I7-2011."),
    ("p", "Schema de legare la pământ este TN-S. Toate masele instalaţiei electrice sunt legate prin conductoare de protecţie la neutrul alimentării legat la pământ (PE)."),
    ("p", "Protecţia împotriva atingerilor indirecte prin întreruperea automată a alimentării se realizează cu dispozitive de protecţie împotriva supracurenţilor. S-a respectat lungimea maximă a buclei de defect. Se prevăd dispozitive de protecţie la curent diferenţial rezidual."),

    ("h2", "SĂNĂTATEA ȘI SECURITATEA MUNCII ÎN TIMPUL EXECUȚIEI"),
    ("p", "Se vor respecta şi aplica toate prevederile de securitate şi sănătate în muncă în vigoare, în scopul asigurării condiţiilor normale de muncă şi evitării accidentelor."),
    ("p", "Coordonarea în materie de securitate şi sănătate trebuie să fie organizată atât în faza de studiu, concepţie şi elaborare a proiectului, cât şi pe perioada executării lucrărilor. Beneficiarul lucrării sau managerul de proiect trebuie să asigure realizarea planului de securitate şi sănătate în muncă care transpune Directiva 89/391/CEE. Pe toată durata realizării lucrării, angajatorul şi lucrătorii independenţi trebuie să respecte obligaţiile generale care le revin în conformitate cu prevederile din legislaţia naţională."),
    ("p", "Cerinţe minime specifice pentru instalaţii electrice:"),
    ("li", "legarea obligatorie la pământ a aparatelor, echipamentelor şi utilajelor care se pot afla în mod accidental sub tensiune;"),
    ("li", "la montajul, punerea în funcţiune, exploatarea şi întreţinerea instalaţiei care face obiectul prezentului proiect se vor respecta normele de tehnica securităţii muncii specifice lucrărilor care se vor executa;"),
    ("li", "alimentarea cu energie electrică a sculelor, echipamentelor şi utilajelor se va face numai de la prize cu contact de protecţie sau tablouri electrice legate la instalaţia de împământare;"),
    ("li", "pentru lucrul la înălţimi mai mari de 2,5m se vor utiliza platforme montate rigid, schele metalice şi centuri de siguranţă;"),
    ("li", "la fiecare loc de muncă vor fi afişate mijloace de avertizare vizuală;"),
    ("li", "dispozitive de protecţie cu chei speciale vor fi montate la uşile tablourilor electrice şi se vor prevedea plăcuţe avertizoare şi alte mijloace pentru interzicerea accesului neautorizat la circuitele electrice;"),
    ("li", "obiectivele proiectate nu se vor pune în funcţiune, parţial sau total, nici măcar pe timp limitat, înainte de asigurarea tuturor măsurilor de tehnica securităţii muncii."),

    ("h2", "SECURITATEA LA INCENDIU ÎN TIMPUL EXECUȚIEI"),
    ("p", "Normele generale de prevenire şi stingere a incendiilor stabilesc principiile, criteriile de performanţă, cerinţele şi condiţiile tehnice privind siguranţa la foc pentru construcţii, instalaţii şi alte amenajări. Normele generale se aplică la proiectarea, executarea şi exploatarea construcţiilor, instalaţiilor şi a altor amenajări, în raport cu faza de realizare în care se află şi indiferent de titularul dreptului de proprietate."),
    ("p", "Cerinţe minime generale specifice instalaţiilor electrice:"),
    ("li", "în caz de incendiu la instalaţiile electrice, înainte de a se acţiona pentru stingerea acestora, se vor scoate de sub tensiune instalaţiile electrice afectate şi cele periclitate;"),
    ("li", "la instalaţiile electrice interioare, pentru stingerea incendiilor se vor folosi numai stingătoare cu praf şi bioxid de carbon;"),
    ("li", "se va asigura verificarea instalaţiilor electrice înainte de punerea sub tensiune;"),
    ("li", "se va asigura utilizarea numai a aparatelor şi echipamentelor electrice aflate în bună stare;"),
    ("li", "se va asigura menţinerea în bună stare a sistemelor de protecţie aferente;"),
    ("li", "se va asigura executarea reparaţiilor, reviziilor şi întreţinerii numai de către personal autorizat;"),
    ("li", "se va asigura preîntâmpinarea acţiunii rozătoarelor asupra învelişului de protecţie din PVC al cablurilor electrice."),
    ("p", "Beneficiarul va lua măsuri ca dotările cu mijloace PSI şi instalaţiile de prevenire şi stingere a incendiilor să fie în perfectă stare de funcţionare."),

    ("h2", "MĂSURI DE PROTECŢIA MUNCII"),
    ("p", "În vederea evitării producerii accidentelor de muncă şi eliminarea pericolelor de electrocutare a personalului în timpul execuţiei şi exploatării instalaţiilor electrice, prin prezentul proiect se prevăd măsuri de protecţia muncii, dintre care cele mai importante sunt:"),
    ("li", "alegerea corespunzătoare a aparatajului în funcţie de mediu şi riscul de incendiu în care acesta funcţionează;"),
    ("li", "amplasarea accesibilă a echipamentelor în vederea unei întreţineri uşoare;"),
    ("li", "prevederea prin proiect a instalaţiei de legare la pământ pentru protecţia împotriva şocurilor electrice;"),
    ("li", "pentru protecţia împotriva atingerilor indirecte, toate elementele metalice ale echipamentelor electrice care în mod normal nu sunt sub tensiune vor fi legate la instalaţia de legare la pământ;"),
    ("li", "dispozitive de protecţie la curent diferenţial rezidual conform paragrafului 2.8."),
    ("p", "Se va acorda o atenţie deosebită următoarelor norme: Legea securităţii şi sănătăţii în muncă nr. 319/2006, Normele metodologice de aplicare aprobate prin HG nr. 1425/2006. Toate lucrările de montaj ale instalaţiilor electrice se vor executa numai de muncitori care au calificarea corespunzătoare şi instructajul de protecţia muncii pentru locul de muncă respectiv."),

    ("h2", "PREVEDERI FINALE"),
    ("p", "Proiectul de instalaţii electrice se verifică de verificator de proiecte atestat conform Legii 10/1995. Beneficiarul va lua toate măsurile necesare respectării prevederilor Legii 10/1995 şi completărilor ulterioare."),
    ("p", "Lucrările vor fi încredinţate spre executare unor firme specializate şi atestate pentru categoriile respective de lucrări."),
    ("p", "Orice modificare intervenită pe parcursul realizării lucrării la execuţie va fi adusă la cunoştinţa proiectantului pentru stabilirea soluţiilor în conformitate cu normativele în vigoare. Efectuarea unor modificări fără avizul proiectantului poate să-l absolve pe acesta de răspunderea faţă de eventualele consecinţe."),

    ("h1", "IV. CERINȚE DE CALITATE ȘI CRITERII DE PERFORMANȚĂ"),
    ("p", "Această exigenţă se apreciază prin:"),
    ("li", "rezistenţa mecanică a elementelor instalaţiei electrice la eforturile exercitate în timpul utilizării;"),
    ("li", "numărul minim de manevre mecanice asupra aparatelor electrice şi asupra corpurilor de iluminat care nu produc deteriorări şi uzură;"),
    ("li", "rezistenţa materialelor, aparatelor şi echipamentelor electrice la condiţii maxime de utilizare;"),
    ("li", "adaptarea măsurilor de protecţie antiseismică;"),
    ("li", "limitarea transmiterii vibraţiilor produse de utilaje şi echipamente electrice susceptibile să intre în rezonanţă."),

    ("h2", "Securitate la incendiu"),
    ("p", "Această exigenţă se apreciază prin:"),
    ("li", "adaptarea instalaţiei electrice la gradul de rezistenţă la foc a elementelor de construcţie;"),
    ("li", "încadrarea instalaţiei electrice în categoriile privind pericolul de incendiu respectiv pericolul de explozie;"),
    ("li", "precizarea nivelului de combustibilitate a componentelor instalaţiei electrice;"),
    ("li", "precizarea limitei de rezistenţă la foc a elementelor de construcţie străpunse de instalaţie."),
    ("p", "Conform normativelor şi standardelor în vigoare se evită montarea instalaţiei electrice pe elemente de construcţie din materiale combustibile. Dacă acest lucru nu este posibil se iau măsuri de protecţie a porţiunii de instalaţie expusă la pericolul de incendiu."),

    ("h2", "Siguranţă în exploatare"),
    ("p", "Această exigenţă se apreciază prin:"),
    ("li", "protecţia utilizatorului împotriva şocurilor electrice prin atingere directă sau indirectă;"),
    ("li", "securitatea instalaţiei electrice la funcţionare în regim anormal (protecţie la suprasarcină, scurtcircuit, scădere de tensiune);"),
    ("li", "limitarea temperaturii exterioare a suprafeţelor accesibile ale echipamentelor electrice;"),
    ("li", "limitarea riscului de rănire prin contact cu părţile în mişcare ale utilajelor şi echipamentelor."),

    ("h2", "Protecţia împotriva zgomotului"),
    ("p", "Această exigenţă se apreciază prin asigurarea confortului acustic în încăperi dotate cu instalaţii electrice ce pot emite zgomote pe perioade scurte de timp, nivelul admis pentru zgomotul emis de instalaţiile electrice din spaţiile tehnice, şi măsurile de limitare a zgomotului în cazul echipamentelor electromagnetice."),

    ("h2", "Igienă, sănătate şi mediu"),
    ("p", "Această exigenţă se apreciază prin evitarea riscului de producere sau favorizare a dezvoltării de substanţe nocive sau insalubre şi limitarea producerii de descărcări electrice care favorizează apariţia şi propagarea incendiului."),

    ("h2", "Economia de energie şi izolare termică"),
    ("p", "Această exigenţă se apreciază prin asigurarea unor consumuri optime de energie electrică, asigurarea unor pierderi minime admise de tensiune, încadrarea consumului de energie activă şi reactivă în limitele admise, şi adoptarea soluţiilor de execuţie care au o valoare minimă a energiei înglobate."),
]


def _memoriu_docx_fv_sections(doc, solar, is_pt=False):
    """G3: capitolele FV în memoriul DOCX — [B] descrierea sistemului (identică pe faze) + [C]
    priza de pământ DEDICATĂ, RAMIFICATĂ pe fază (fix Dan): PT -> breviarul complet pe pattern-ul
    M5 (_brevier_du_block: 'X = descriere = valoare'); DTAC -> DOAR fraza de propunere (fără
    cifre). Aceleași date și cazuri speciale ca memoriul text (G2, main._memoriu_fv_sections).
    `solar` gol/absent -> nimic (backward-compat); fail-safe: orice eroare -> fără capitolele FV."""
    if not isinstance(solar, dict) or not (solar.get("package_kw") or solar.get("power_kw")):
        return
    try:
        import math as _m
        from schema_fv import FV_PACKAGES, FV_FIXED, snap_fv_package, fv_grounding
        kw = snap_fv_package(solar.get("package_kw") or solar.get("power_kw"))
        pkg = FV_PACKAGES[kw]
        g = fv_grounding(kw, solar.get("soil_type"))
        soil_lbl = {"mlastinos": "sol mlăștinos", "argila": "argilă umedă", "agricol": "sol agricol",
                    "nisip_umed": "nisip umed", "nisip_uscat": "nisip uscat", "pietris": "pietriș"}[g["soil_type"]]

        # [B] 2.8 + 2.8.1 descrierea sistemului (montajul = "conform planului": memoriul se
        # generează înaintea plasării tablourilor). Subsecțiunile ca paragrafe bold (deciziile Dan).
        _add_heading(doc, "2.8. SISTEM FOTOVOLTAIC", level=2)
        _add_para(doc, "2.8.1. Descrierea sistemului", bold=True)
        nstr = "{} string-uri".format(pkg["nr_stringuri"]) if pkg["nr_stringuri"] != 1 else "1 string"
        _add_para(doc,
                  "Obiectivul este echipat cu un sistem fotovoltaic format din {} panouri fotovoltaice "
                  "monocristaline de {} Wp fiecare, cu o putere instalată totală de {:.2f} kWp, dispuse "
                  "în {}. Conversia se realizează printr-un invertor solar trifazat de {} kW, montat pe "
                  "fațadă sau în spațiul tehnic, conform planului de forță. Sistemul se racordează la "
                  "tabloul electric general (TEG) prin tablourile de interfață și protecție T.CC (curent "
                  "continuu) și T.CA (curent alternativ). Producția este contorizată separat prin contor "
                  "de producție 400V.".format(pkg["nr_panouri"], FV_FIXED["wp"], pkg["pi_kw"], nstr,
                                              pkg["invertor_ca_kw"]))
        _blank(doc, 1)

        # [C] 2.8.2 priza de pământ dedicată — RAMURA PE FAZĂ: DTAC = fraza de propunere;
        # PT = breviarul complet (pattern M5: rânduri 'X = descriere = valoare')
        _add_para(doc, "2.8.2. Priza de pământ a sistemului fotovoltaic", bold=True)
        _add_para(doc, "Sistemul fotovoltaic se leagă la o priză de pământ DEDICATĂ, separată de "
                       "priza de pământ a instalației generale.")
        if not is_pt:
            _add_para(doc, "Pentru sistemul fotovoltaic se va propune o priză de pământ dedicată, "
                           "dimensionată la faza PT (Rp ≤ 4 Ω).")
            _blank(doc, 1)
            return
        _add_para(doc, "Ipoteze de calcul:", bold=True)
        _add_para(doc, "ρ = rezistivitatea solului = {} Ω·m ({})".format(g["rho"], soil_lbl))
        _add_para(doc, "Rp = rezistența țintă ≤ 4 Ω (I7-2011)")
        _add_para(doc, "Țăruși OL-Zn Ø18 mm, L = 1,5 m, distanța între țăruși a = 3 m (a/L = 2)")
        _add_para(doc, "Priză de contur: țăruși legați cu platbandă OL-Zn 40x4 mm, îngropată la 0,8 m")
        r1 = g["rho"] / (2 * _m.pi * 1.5) * (_m.log(4 * 1.5 / 0.018) - 1.0)
        _add_para(doc, "Rezistența de dispersie a unui țăruș:", bold=True)
        _add_para(doc, "R1 = ρ/(2·π·L)·[ln(4L/d) − 1] = {:.1f} Ω".format(r1))
        _add_para(doc, "Rezistența grupului de n țăruși:", bold=True)
        _add_para(doc, "Rn = R1/(n·η), cu η ≈ 0,7 (coeficientul de utilizare la a/L = 2), "
                       "în paralel cu platbanda de contur")
        _add_para(doc, "Dimensionare ({}, sistem de {} kW): {} țăruși de 1,5 m + {} m platbandă "
                       "OL-Zn 40x4".format(soil_lbl, kw, g["tarusi"], g["platbanda_m"]), bold=True)
        if g["soil_type"] == "pietris":
            # nota SPECIALĂ pietriș: înlocuiește linia de Rp (țărușii de 1,5 m nu ating ținta)
            _add_para(doc, "NOTĂ: În sol pietros (ρ = 500 Ω·m), țărușii de 1,5 m nu asigură Rp ≤ 4 Ω. "
                           "Se recomandă țăruși de 2-3 m, electrozi forați sau tratarea solului "
                           "(bentonită). Rp se măsoară la recepție.")
        else:
            rp = float(g["rp"])
            _add_para(doc, "Rp = rezistența de dispersie estimată = {} Ω{}".format(
                g["rp"], " ≤ 4 Ω" if rp <= 4.0 else ""))
            if rp > 4.0:
                _add_para(doc, "NOTĂ: Rp estimat depășește 4 Ω; se vor adăuga țăruși suplimentari sau "
                               "se prelungește conturul până la atingerea valorii ≤ 4 Ω. Rp se măsoară "
                               "la recepție.")
        _add_para(doc, "Rezistența se măsoară la recepție (metoda Wenner). Valorile sunt estimate de calcul.")
        _blank(doc, 1)
    except Exception:
        pass


def _page_memoriu(doc, cp, cf, solar=None):
    _add_heading(doc, "III. MEMORIU TEHNIC INSTALAȚII ELECTRICE", level=1)
    _add_heading(doc, "1. DATE GENERALE", level=2)

    _add_label_value(doc, "1.1. Denumirea lucrării: ",
                     cp.get("titlu_proiect", ""))
    _add_label_value(doc, "1.2. Adresa: ", cp.get("amplasament", ""))
    _add_label_value(
        doc,
        "1.3. Obiect: ",
        "Prezentul memoriu tehnic descrie soluțiile tehnice adoptate pentru "
        "realizarea instalațiilor electrice aferente obiectivului menționat "
        "mai sus.",
    )

    # Text fix sectiuni tehnice — kind in {"h1","h2","p","li"}.
    # G3 (repozitionare Dan): cu FV selectat, sectiunea FV intra la 2.8 (dupa 2.7 Iluminat),
    # iar "2.8. PROTECTIA..." devine "2.9. ..." + referinta incrucisata "paragrafului 2.8." -> 2.9.
    # Match pe STRINGURI COMPLETE (nu "2.8" izolat — "art. 4.2.2.8 din I7-2011" ramane NEATINS).
    # Fara FV: blocurile ies exact ca inainte (non-regresie).
    _has_fv = isinstance(solar, dict) and bool(solar.get("package_kw") or solar.get("power_kw"))
    for kind, text in _MEMORIU_BLOCKS:
        if _has_fv and kind == "h2" and text.startswith("2.8. PROTEC"):
            _memoriu_docx_fv_sections(doc, solar, is_pt=_is_pt(cp.get("faza")))   # 2.8 = FV
            text = "2.9. " + text[len("2.8. "):]                                  # PROTECTIA -> 2.9
        elif _has_fv and kind == "li" and "paragrafului 2.8." in text:
            text = text.replace("paragrafului 2.8.", "paragrafului 2.9.")
        if kind == "h1":
            _add_heading(doc, text, level=1)
        elif kind == "h2":
            _add_heading(doc, text, level=2)
        elif kind == "li":
            _add_para(doc, "- " + text)
        else:
            _add_para(doc, text)

    # Semnatura finala in tabel cu chenar (dinamic din cartus_firma), centrat
    doc.add_paragraph()
    sig_left_w = Cm(4.5)
    sig_right_w = Cm(4.0)
    sig_rows = [
        ("PROIECTANT DE SPECIALITATE", cf.get("firma_nume", "")),
        ("ÎNTOCMIT", cf.get("proiectant_nume", "")),
    ]
    sig = doc.add_table(rows=0, cols=2)
    sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig.allow_autofit = False
    _set_table_fixed(sig)
    for col, w in zip(sig.columns, (sig_left_w, sig_right_w)):
        col.width = w
    for label, value in sig_rows:
        cells = sig.add_row().cells
        c_left, c_right = cells[0], cells[1]
        c_left.width = sig_left_w
        c_right.width = sig_right_w
        _set_run_font(c_left.paragraphs[0].add_run(label), size=11, bold=True)
        _set_run_font(c_right.paragraphs[0].add_run(value), size=11, bold=False)
    set_table_borders(sig)  # chenar vizibil pe tabelul de semnatura


# =============================================================================
# V. BREVIAR DE CALCUL  (M5-B) — DOAR la PT (gate is_pt)
# =============================================================================

_SQRT3 = 3 ** 0.5

# Tabel selecție cablu (curent_max_A, secţiune_mmp) — replică locală a CABLE_SECTIONS
# din main.py. Păstrat self-contained: memoriu_generator.py NU importă app-ul FastAPI.
_CABLE_SECTIONS = [
    (6, "1.5"), (10, "2.5"), (16, "4"), (25, "6"),
    (32, "10"), (40, "16"), (63, "25"), (float("inf"), "35"),
]


def _cable_for_current(current_a):
    """Secţiunea standard de cupru (mmp) al cărei Iz acoperă curentul dat.
    Convenţie (ca în main.py): se pasează curentul de calcul deja înmulţit cu 1,25."""
    try:
        ia = float(current_a)
    except (TypeError, ValueError):
        return "1.5"
    for limit, section in _CABLE_SECTIONS:
        if ia <= limit:
            return section
    return "35"


def _sect_ro(section):
    """Formatare secţiune pentru afişare RO (virgulă zecimală): '2.5' -> '2,5'."""
    return str(section).replace(".", ",")


def _du_trifazat_pct(ic_a, l_m, s_mmp, cosphi=0.92, un_v=400.0, gamma=59.6):
    """Cădere de tensiune procentuală pe o coloană trifazată:
        ΔU% = (√3 · Ic · L · cosφ · 100) / (γ · S · Un)
    Formula standard (verificată pe exemplul TE-CT din PT: 22,95 A / 20 m / 4 mmp -> 0,77%).
    Notă: exemplul TEG din PT (65,14 A / 40 m / S=10) afişează 1,24%, dar 1,24% corespunde
    de fapt lui S=14 mmp — sursa e inconsistentă pe secţiune; formula de aici e cea corectă fizic.
    Robustă: secţiune/tensiune/gamma invalide -> 0.0 (nu crapă)."""
    try:
        s = float(s_mmp)
        ic = float(ic_a)
        l = float(l_m)
    except (TypeError, ValueError):
        return 0.0
    if s <= 0 or un_v <= 0 or gamma <= 0:
        return 0.0
    return (_SQRT3 * ic * l * float(cosphi) * 100.0) / (gamma * s * un_v)


# Formule standard pentru curenţii nominali (transcrise din exemplul PT — IE.00 MT.ELECTRICE).
_BREVIAR_FORMULE_IN = [
    "In = Pi/(Uf·cosφ), pentru circuite monofazate de lumină, în care Pi este puterea instalată în waţi, Uf=220V, cosφ=1 pentru lămpi incandescente şi 0,95 pentru lămpi fluorescente;",
    "In = Pi/(Uf·cosφ·η), pentru circuite monofazate de prize, în care Pi este puterea instalată în waţi, Uf=220V; η (randamentul) se consideră 0,85 pentru diferite receptoare introduse în priză;",
    "In = Cc·Pi/(1,73·Ul·cosφ), pentru coloanele secundare trifazate echilibrate ale tablourilor principale de nivel, în care Pi este puterea instalată în waţi, Ul=400V, Cc este coeficientul de cerere pe coloană, cosφ = factorul de putere mediu calculat al coloanei;",
    "In = Cc·Pi/(Uf·cosφ), pentru coloanele secundare trifazate dezechilibrate (încărcate asimetric), calculul făcându-se pe faza cea mai încărcată, Uf = tensiunea de fază în volţi, ceilalţi factori având aceeaşi semnificaţie;",
    "In = Cc·Pi/(1,73·Ul·cosφ), pentru coloanele principale, inclusiv coloana principală a TEG, în care Pi = suma puterilor instalate pe coloanele secundare, Cc = factorul de cerere al coloanei principale, ceilalţi factori având aceeaşi semnificaţie.",
]

# Formule standard pentru verificarea la cădere de tensiune (transcrise din exemplul PT).
_BREVIAR_FORMULE_DU = [
    "ΔU% = [2·100·Σ(Cc·Pi·li)/Si]/(γ·Uf), pentru circuite şi coloane monofazate cu sarcini uniform distribuite.",
    "ΔU% = [100·Σ(Cc·Pi·li)/Si]/(γ·Ul), pentru circuite trifazate cu mai multe receptoare concentrate.",
]


def _brevier_du_block(doc, titlu, ic_a, l_m, s_ro, du_pct):
    """Un bloc 'Tablou X' din secţiunea Căderea de tensiune (TEG / TE-CT)."""
    _add_para(doc, titlu, bold=True)
    _add_para(doc, "Un = tensiunea de alimentare = 400 V")
    _add_para(doc, "Ic = curentul electric de calcul = {:.2f} A".format(ic_a))
    _add_para(doc, "L = lungimea cablului = {:.0f} m".format(l_m))
    _add_para(doc, "cosφ = factorul de putere = 0,92")
    _add_para(doc, "S = secţiunea cablului = {} mmp".format(s_ro))
    _add_para(doc, "γ = conductibilitate cupru = 59,6 Ω·mmp/m")
    verdict = "< 5%" if du_pct < 5.0 else "≥ 5% (ATENȚIE: depăşeşte limita admisă!)"
    _add_para(doc, "ΔU = {:.2f}% {}".format(du_pct, verdict))
    _blank(doc, 1)


def _page_brevier(doc, cp, cf, circuits, power_summary):
    """V. BREVIAR DE CALCUL — DOAR la PT (gate is_pt). Formule standard (transcrise din
    exemplul PT) + exemple lucrate cu numere REALE din circuits/power_summary (M5-A):
      - Lumină: circuitul de iluminat cel mai încărcat (power_w maxim) -> Pi real.
      - Prize : 2000 W FIX (standard, nu se schimbă).
      - TEG   : Ic real (power_summary.current_a), L=20 m fix, S din main_breaker_a, ΔU calculat.
      - TE-CT : valori standard reprezentative (nu există date TE-CT în payload-ul memoriului).
    Backward-compat: fără circuits/power_summary -> valori standard (Pi=360W, Ic=65,14A), NU crapă."""
    # Coerciție pe TIP (nu pe truthiness): un payload truthy dar de tip greșit
    # (power_summary listă/string, circuits scalar) nu trebuie să crape build-ul.
    circuits = circuits if isinstance(circuits, list) else []
    power_summary = power_summary if isinstance(power_summary, dict) else {}

    # --- Lumină: circuitul de iluminat cel mai încărcat (power_w maxim) ---
    lum_w = 0.0
    for c in circuits:
        if not isinstance(c, dict):
            continue
        if str(c.get("type", "")).lower().startswith("iluminat"):
            try:
                p = float(c.get("power_w", 0) or 0)
            except (TypeError, ValueError):
                p = 0.0
            if p > lum_w:
                lum_w = p
    if lum_w <= 0:
        lum_w = 360.0                              # backward-compat: standard ca în exemplu
    ic_lum = lum_w / 230.0                          # Ic = Pi/U (cosφ=1, lumină)
    s_lum = _cable_for_current(ic_lum * 1.25)       # secţiune (curent de calcul ×1,25)

    # --- TEG: Ic real + S derivat din main breaker ---
    try:
        ic_teg = float(power_summary.get("current_a", 0) or 0)
    except (TypeError, ValueError):
        ic_teg = 0.0
    if ic_teg <= 0:
        ic_teg = 65.14                              # backward-compat: reprezentativ (exemplu)
    try:
        mb = float(power_summary.get("main_breaker_a", 0) or 0)
    except (TypeError, ValueError):
        mb = 0.0
    s_teg = _cable_for_current(mb if mb > 0 else ic_teg * 1.25)
    l_teg = 20.0                                    # L = 20 m FIX (decizie Dan; intenţionat NU 40 m din exemplu)
    du_teg = _du_trifazat_pct(ic_teg, l_teg, s_teg)

    # --- TE-CT: standard reprezentativ (nu există date TE-CT în payload-ul memoriului) ---
    ic_tect, l_tect, s_tect = 22.95, 20.0, "4"
    du_tect = _du_trifazat_pct(ic_tect, l_tect, s_tect)

    # --- Render ---
    doc.add_page_break()
    _add_heading(doc, "V. BREVIAR DE CALCUL", level=1)
    _add_para(doc, "BREVIAR DE CALCUL INSTALAȚII ELECTRICE", bold=True)
    _add_para(doc, "Determinarea secţiunii circuitelor şi coloanelor de alimentare.")
    _add_para(doc, "Determinarea curenţilor absorbiţi (nominali) de circuite şi coloane s-a făcut "
                   "utilizând următoarele formule de calcul:")
    for f in _BREVIAR_FORMULE_IN:
        _add_para(doc, f)
    _add_para(doc, "După calculul secţiunilor circuitelor şi coloanelor, acestea se verifică la "
                   "pierderile de tensiune. S-au utilizat următoarele formule de calcul:")
    for f in _BREVIAR_FORMULE_DU:
        _add_para(doc, f)
    _add_para(doc, "Datorită distanțelor relativ mici între tabloul TEG și consumatori și a "
                   "gradului mic de încărcare a circuitelor nu se pune problema unor căderi de "
                   "tensiune inacceptabile.")
    _blank(doc, 1)

    # 1. Dimensionarea conductelor electrice
    _add_para(doc, "1. Dimensionarea conductelor electrice:", bold=True)
    _add_para(doc, "Pe circuitul de lumină:", bold=True)
    _add_para(doc, "Ic = Pi : U : cosφ  [A]")
    _add_para(doc, "Pi = Puterea instalată")
    _add_para(doc, "U = Tensiunea de alimentare")
    _add_para(doc, "cosφ = Factorul de putere")
    _add_para(doc, "Ic = Curentul de calcul")
    _add_para(doc, "Se va dimensiona pentru circuitul cel mai încărcat: Pi = {:.0f} W".format(lum_w))
    _add_para(doc, "Ic = {:.0f} : 230".format(lum_w))
    _add_para(doc, "Ic = {:.2f} A".format(ic_lum))
    _add_para(doc, "Se alege conductor din cupru {} mmp".format(_sect_ro(s_lum)))
    _blank(doc, 1)

    # 2. Alegerea siguranțelor automate
    _add_para(doc, "2. Alegerea siguranțelor automate", bold=True)
    _add_para(doc, "• Pe circuitul de iluminat:")
    _add_para(doc, "If < k · Imax")
    _add_para(doc, "Imax = curentul admis (pentru conductor de Cu de 1,5 mmp = 14 A)")
    _add_para(doc, "k = coeficientul de siguranță = 0,8")
    _add_para(doc, "If < 0,8 · 14")
    _add_para(doc, "If < 11,2")
    _add_para(doc, "Se alege siguranță automată 1P+N, 10A")
    _add_para(doc, "Pe circuitul de prize:")
    _add_para(doc, "Ic = 2000 : 230")
    _add_para(doc, "Ic = 8,69 A")
    _add_para(doc, "Se alege conductor din cupru 2,5 mmp")
    _blank(doc, 1)

    # Căderea de tensiune (TEG real + TE-CT standard)
    _add_para(doc, "Căderea de tensiune:", bold=True)
    _brevier_du_block(doc, "Tablou TEG", ic_teg, l_teg, _sect_ro(s_teg), du_teg)
    _brevier_du_block(doc, "Tablou TE-CT", ic_tect, l_tect, _sect_ro(s_tect), du_tect)


# Faze determinante STANDARD pentru instalaţii electrice (verificări obligatorii — din exemplul PT).
_FAZE_DETERMINANTE = [
    "Verificarea rezistenţei la dispersie a prizei de pământ existenta.",
    "Verificarea rezistenţei la dispersie a prizei de pământ propusa.",
    "Verificarea legării la pământ a instalației electrice.",
    "Verificarea legarii la pamant a tuturor maselor metalice.",
]


def _page_faze_determinante(doc, cp, cf):
    """VI. FAZE DETERMINANTE — DOAR la PT (gate is_pt). Antet din cartus_proiect (cp) +
    lista standard de verificări + semnături. Stil DTAC (helperele existente)."""
    doc.add_page_break()
    _add_heading(doc, "VI. FAZE DETERMINANTE PENTRU INSTALAȚII ELECTRICE", level=1)
    _add_label_value(doc, "BENEFICIAR: ", cp.get("beneficiar", ""))
    _add_label_value(doc, "LUCRARE: ", cp.get("titlu_proiect", ""))
    _add_label_value(doc, "ADRESA: ", cp.get("amplasament", ""))
    _add_label_value(doc, "FAZA: ", cp.get("faza", ""))
    _blank(doc, 1)
    for f in _FAZE_DETERMINANTE:
        _add_para(doc, "- " + f)
    _blank(doc, 1)
    _add_para(doc, "Inspector de specialitate (numele și prenumele) …………………………………………")
    _add_para(doc, "Semnătura/ștampila …………………………………………………………………………")
    _blank(doc, 1)
    _add_centered(doc, "Proiectant de specialitate", size=11, bold=True)
    _add_centered(doc, cf.get("firma_nume", ""), size=11, bold=False)
    _add_centered(doc, cf.get("proiectant_nume", ""), size=11, bold=False)


# =============================================================================
# VII. PROGRAM DE CONTROL AL CALITĂȚII  (M4) — DOAR la PT (gate is_pt)
# =============================================================================

# Antet coloane + cele 11 rânduri de control (transcrise fidel din exemplul PT — IE.00
# MT.ELECTRICE). Conţinutul (cuvintele) e verbatim; s-au normalizat doar salturile de rând
# din interiorul celulelor (un element pe linie) şi diacriticele evident pierdute.
_PROGRAM_CONTROL_HEADER = [
    "Nr. crt.",
    "Lucrări ce se controlează, verifică sau se recepționează",
    "Documentul scris care se încheie",
    "Cine participă",
    "Nr. / data actului încheiat",
]

_PROGRAM_CONTROL_ROWS = [
    ("1", "Verificarea proiectului de instalații electrice de verificatori de proiecte atestați de M.L.P.A.T.",
     "Referat verificare", "Verificator Atestat\nBeneficiar\nProiectant", ""),
    ("2", "Predarea proiectului executantului",
     "Proces verbal", "Beneficiar\nExecutant", ""),
    ("3", "Verificarea calității materialelor utilizate și a echipamentelor procurate",
     "Proces verbal", "Beneficiar\nExecutant", ""),
    ("4", "Verificare priză de pământ",
     "Proces verbal", "Beneficiar\nExecutant", ""),
    ("5", "Execuție trasee circuite electrice",
     "Proces verbal", "Beneficiar\nExecutant", ""),
    ("6", "Montare aparate electrice, corpuri de iluminat, tablouri electrice, echipament electric",
     "Proces verbal", "Beneficiar\nExecutant", ""),
    ("7", "Încercarea continuității electrice a circuitelor electrice\nVerificarea corpurilor de iluminat\n"
          "Încercarea aparatelor electrice\nÎncercarea tablourilor electrice",
     "Proces verbal", "Beneficiar\nExecutant\nProiectant\nFurnizori", ""),
    ("8", "FD: Verificarea rezistenței de dispersie a prizei de pământ\n"
          "Verificarea legării la pământ a instalației electrice",
     "Proces verbal\nBuletin de verificare priză de pământ", "Beneficiar\nExecutant\nProiectant", ""),
    ("9", "Recepția la terminarea lucrărilor",
     "Proces verbal de constatare funcționării", "Comisia de recepție", ""),
    ("10", "Urmărirea calității și funcționării instalațiilor",
     "", "Beneficiar", ""),
    ("11", "Recepția finală",
     "", "Comisia de recepție", ""),
]

# Lăţimi coloane (total ~16,5 cm = lăţimea utilă A4 cu marginile din _setup_document).
_PROGRAM_CONTROL_WIDTHS = (Cm(1.2), Cm(6.6), Cm(3.3), Cm(3.2), Cm(2.2))


def _program_cell(cell, text, width, bold=False, bg=None, size=10):
    """Scrie o celulă de tabel (multi-linie pe '\\n'), font Arial, opţional bold/fundal."""
    cell.width = width
    if bg:
        set_cell_bg(cell, bg)
    lines = str(text).split("\n")
    p = cell.paragraphs[0]
    _set_run_font(p.add_run(lines[0]), size=size, bold=bold)
    for ln in lines[1:]:
        p2 = cell.add_paragraph()
        _set_run_font(p2.add_run(ln), size=size, bold=bold)


def _page_program_control(doc, cp, cf):
    """VII. PROGRAM DE CONTROL AL CALITĂȚII — DOAR la PT (gate is_pt). ULTIMA secţiune a
    memoriului PT. Antet din cartus_proiect (cp, ca la VI. Faze determinante) + tabel 12×5
    (antet + 11 rânduri, transcris fidel din exemplul PT) + semnătură proiectant din cf.
    Stil DTAC (helperele existente: _set_table_fixed / set_cell_bg / set_table_borders)."""
    doc.add_page_break()
    _add_heading(doc, "VII. PROGRAM DE CONTROL AL CALITĂȚII LUCRĂRILOR DE INSTALAȚII ELECTRICE", level=1)
    _add_label_value(doc, "BENEFICIAR: ", cp.get("beneficiar", ""))
    _add_label_value(doc, "LUCRARE: ", cp.get("titlu_proiect", ""))
    _add_label_value(doc, "ADRESA: ", cp.get("amplasament", ""))
    _add_label_value(doc, "FAZA: ", cp.get("faza", ""))
    _blank(doc, 1)

    table = doc.add_table(rows=0, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    _set_table_fixed(table)
    for col, w in zip(table.columns, _PROGRAM_CONTROL_WIDTHS):
        col.width = w

    # antet (gri + bold)
    hdr = table.add_row().cells
    for cell, txt, w in zip(hdr, _PROGRAM_CONTROL_HEADER, _PROGRAM_CONTROL_WIDTHS):
        _program_cell(cell, txt, w, bold=True, bg=CELL_BG_GREY)
    # rânduri de control
    for row in _PROGRAM_CONTROL_ROWS:
        cells = table.add_row().cells
        for cell, txt, w in zip(cells, row, _PROGRAM_CONTROL_WIDTHS):
            _program_cell(cell, txt, w)

    set_table_borders(table)  # chenare vizibile (stil DTAC)

    _blank(doc, 1)
    _add_centered(doc, "Proiectant de specialitate", size=11, bold=True)
    _add_centered(doc, cf.get("firma_nume", ""), size=11, bold=False)
    _add_centered(doc, cf.get("proiectant_nume", ""), size=11, bold=False)


# =============================================================================
# ENTRYPOINT
# =============================================================================

def build_memoriu_docx(data: dict) -> bytes:
    """Construieste memoriul .docx din datele cartusului si returneaza bytes."""
    data = data or {}
    cp = data.get("cartus_proiect") or {}
    cf = data.get("cartus_firma") or {}
    planse = data.get("planse") or []
    circuits = data.get("circuits") or []          # M5-A: circuite (type/power_w) -> brevier
    power_summary = data.get("power_summary") or {} # M5-A: current_a (Ic TEG), main_breaker_a
    solar = data.get("solar") or {}                # G3: pachetul FV + solul (n8n; gol -> fara capitole FV)
    is_pt = _is_pt(cp.get("faza"))   # PT -> borderou extins + secțiuni noi (M2-M5); DTAC -> NESCHIMBAT

    # PT: lista REALĂ de planșe (planuri per nivel + scheme monofilare care EXISTĂ), numerotată
    # IE.1..IE.N secvențial FĂRĂ goluri (plansa_numbering) — înlocuiește vechea listă hardcodată
    # IE.1..IE.8. DOAR secțiunea de planșe se schimbă; restul memoriului = NEATINS. extra_floors/
    # has_tect: explicit din payload (n8n Faza 2B), altfel derivat din circuite (fallback). Orice
    # eroare -> păstrează `planse` primit / lista standard (backward-compat).
    if is_pt:
        try:
            from plansa_numbering import compute_plansa_numbering, derive_extra_floors
            _has_tect = data.get("has_tect")
            if _has_tect is None:
                _has_tect = any((c or {}).get("panel") == "TE-CT" for c in circuits)
            _extra = data.get("extra_floors")
            if _extra is None:
                _extra = derive_extra_floors(circuits)
            # G3: cu FV selectat (solar prezent), borderoul include si plansa FV (ultima IE)
            _real = compute_plansa_numbering(_extra, bool(_has_tect),
                                             has_fv=bool(solar.get("package_kw") or solar.get("power_kw")))
            if _real:
                planse = [{"nr": p["nr"], "titlu": p["nume"]} for p in _real]
        except Exception:
            pass

    doc = _setup_document()
    _page_coperta(doc, cp, cf)
    _page_fisa(doc, cp, cf)
    _page_borderou(doc, planse, is_pt=is_pt)
    _page_memoriu(doc, cp, cf, solar=solar)
    if is_pt:
        # secțiuni NOI de PT (V. Brevier = M5-B; VI. Faze determinante = M3; VII. Program = M4)
        _page_brevier(doc, cp, cf, circuits, power_summary)
        _page_faze_determinante(doc, cp, cf)
        _page_program_control(doc, cp, cf)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
